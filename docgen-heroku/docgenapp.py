from flask import Flask, request
from bson.objectid import ObjectId
from flask_cors import CORS
from _io import BytesIO, StringIO
import os
from base64 import b64decode
import re
import json
from datetime import date
import datetime
import copy
import base64
import io
from decouple import config
import traceback


# importing model classes
from Models.objwrapper import obj_wrap
from Models.fieldwrap import field_wrap_obj
from Models.grandwrap import grand_wrap_obj
from Models.parentwrap import parent_wrap_obj
from Models.childwrap import child_wrap_obj

# importing dao classes
from Doa.mongo_connect import logging_mongo
from Doa.box_connect import logging_box

# Document reading and writing library
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT

# Utils classes
import query_formation
import utility

# The locale module opens access to the POSIX locale database and functionality
import locale
# user’s preferred locale settings
locale.setlocale(locale.LC_ALL, 'en_US.utf-8')


# Connecting Box.com
boxClient = logging_box().connect_box()
# Connecting MongoDB
db = logging_mongo().connect_mongo()

# insert file to the box.com using the folder id and local file path
# print(logging_box().insert_file(boxClient, '123484346174',
#                                 '/Users/gowtham_kalaiselvan/Documents/Aspigrow/mongopython/Templates/Order Form Child New3.docx'))

app = Flask(__name__)
CORS(app, resources={"/*": {"origins": "*"}})

app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
file_path = ''
document_data = ''
document_data_initial = ''
record_id = ''
file_name = ''
file_type = ''

table_pattern_list = []

folder_id_dyn = ''

pick_list_fields = []
pick_list_fields_in_child_obj = []

port_number = config('PORT')
print('Document Generation Application Listens in PORT -->  ', port_number)


@ app.route('/generate_file', methods=['POST'])
def generate_document():
    try:
        # getting file id and record id from headers from request headers
        file_id = request.headers['file_id']
        record_id = request.headers['record_id']

        # getting file content from Box.com using file id
        file_content = logging_box().get_file_content(boxClient, file_id)

        file_name = 'test'+'.docx'

        # converting file contents to source stream
        bytes = b64decode(file_content)
        source_stream = BytesIO(file_content)
        doc = Document(source_stream)
        source_stream.close()

        # Initialize classes
        generate_mongo_query = query_formation.generate_mongo_query()
        utility_class = utility.utility()

        full_text = []
        field_list = []
        child_obj_metadata = []

        for para in doc.paragraphs:
            full_text.append(para.text)

        document_data = '\n'.join(full_text)
        field_list = re.findall("\\$\\{(.*?)\\}", document_data)
        child_obj_metadata = re.findall("\\$tbl\\{(.*?)\\}", document_data)
        withouttable = utility_class.get_all_table_patterns(
            document_data.replace('\n', ' ').replace('\r', ''), table_pattern_list)
        if len(withouttable) > 0:
            child_obj_metadata = re.findall("\\$\\{(.*?)\\}", withouttable[0])
            field_list = list(set(field_list) - set(child_obj_metadata))
        head_child_obj = ''
        alltext_in_tbl = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    alltext_in_tbl.append(cell.text)
        alltext_in_tbl = '\n'.join(alltext_in_tbl)
        check_value_str = re.findall("\\$tbl\\{(.*?)\\}", alltext_in_tbl)
        table_patterns = re.findall(
            "\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:.*?\\}", alltext_in_tbl.replace('\n', ' ').replace('\r', ''))
        check_is_child_obj = 'Nil'
        child_table_list = []
        child_table_pattern_list = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
                    fields_in_cell = re.findall("\\$\\{(.*?)\\}", cell.text)
                    child_obj_values = re.findall(
                        "\\$tbl\\{(.*?)\\}", cell.text)
                    child_obj_str = []
                    if len(child_obj_values) > 0:
                        child_obj_str = re.findall(
                            "\\$tbl\\{START:(.*?)\\}", cell.text)
                        is_child_table = child_obj_values[0].split(':')
                        obj_name = is_child_table[1].split('.')
                        check_is_child_obj = '#' + \
                            obj_name[0].split(
                                '#')[1]+'#'+obj_name[1].split('(')[0]
                        if check_is_child_obj not in child_table_list:
                            child_table_list.append(check_is_child_obj)
                            child_table_pattern_list.append(
                                child_obj_str[0].split('.')[1])
                    if len(child_obj_str) > 0:
                        head_child_obj = child_obj_str[0]
                    if len(fields_in_cell) > 0:
                        for field in fields_in_cell:
                            if not re.search(check_is_child_obj, field):
                                field_list.append(field)
                            else:
                                child_obj_metadata.append(field.split(':')[1])
        field_list = list(dict.fromkeys(field_list))
        if not len(table_patterns) > 0 and not len(field_list) > 0:
            doc.save('test.docx')
            docx_stream = io.BytesIO()
            doc.save(docx_stream)
            docx_bytes = docx_stream.getvalue()
            encoded = base64.b64encode(docx_bytes)
            doc_data = {"isSuccess": True, "data": {
                "file_id": file_id,
                "file": str(encoded)[2:-1],
                "record_id": record_id,
                "success": True
            }}
            return json.dumps(doc_data)

        obj_wrapper = type('test', (object,), {})()

        # This block will do json formation of fields and parent objects which will be returned to salesforce for retrieving data
        if len(field_list) > 0:
            obj_wrapper = obj_wrap(field_list[0].split('.')[
                                   0], False, [], [], [])
            field_wrapper = []
            parent_wrapper = []
            parent_field_wrapper = []
            grand_parent_field_wrapper = []
            grand_wrapper = []
            for field in field_list:
                date_type = re.findall("( #DATE.*)", field)
                format_type = re.findall("( #[A-Z]*)", field)
                if len(format_type) > 0 and format_type[0] == ' #PICKLIST':
                    corrected_field = field.replace(
                        format_type[0].strip(), '').strip()
                    parent_obj = corrected_field.split('.')
                    pick_list_fields.append(parent_obj[-1])
                elif len(date_type) > 0:
                    corrected_field = field.replace(
                        date_type[0].strip(), '').strip()
                elif len(format_type) > 0:
                    corrected_field = field.replace(
                        format_type[0].strip(), '').strip()
                else:
                    corrected_field = field

                parent_obj = corrected_field.split('.')
                if len(parent_obj) == 2:
                    field_wrap = field_wrap_obj(parent_obj[-1], False)
                    if field_wrap not in field_wrapper:
                        field_wrapper.append(field_wrap)

                else:
                    field_wrap = field_wrap_obj(parent_obj[1], False)
                    if field_wrap.__dict__ not in field_wrapper:
                        field_wrapper.append(field_wrap.__dict__)
                    filtered_obj = parent_obj[1:len(parent_obj)]

                    if len(filtered_obj) == 3:
                        parent_field_wrap = field_wrap_obj(
                            filtered_obj[1], False)
                        if parent_field_wrap not in parent_field_wrapper:
                            parent_field_wrapper.append(parent_field_wrap)
                        grand_parent_field_wrap = field_wrap_obj(
                            filtered_obj[2], False)
                        if grand_parent_field_wrap not in grand_parent_field_wrapper:
                            grand_parent_field_wrapper.append(
                                grand_parent_field_wrap)
                        grand_wrap = grand_wrap_obj(
                            filtered_obj[1], False, grand_parent_field_wrapper)
                        if len(grand_wrapper) > 0:
                            check_grobj_list = list()
                            for obj in grand_wrapper:
                                check_grobj_list.append(obj.objName)

                            if grand_wrap.objName not in check_grobj_list:
                                grand_wrapper.append(grand_wrap)

                            elif {
                                    'fieldName': filtered_obj[2]
                            } not in grand_wrapper[check_grobj_list.index(
                                    grand_wrap.objName)].fieldWrapperList:
                                grand_wrapper[check_grobj_list.index(
                                    grand_wrap.objName)].fieldWrapperList.append(
                                        field_wrap_obj(filtered_obj[2], False))
                        else:
                            grand_wrapper.append(grand_wrap)
                        parent_wrap = parent_wrap_obj(
                            filtered_obj[0], False, parent_field_wrapper, [], [], grand_wrapper)

                        if len(parent_wrapper) > 0:
                            check_obj_list = list()
                            for obj in parent_wrapper:
                                check_obj_list.append(obj.objName)
                            if parent_wrap.objName not in check_obj_list:
                                parent_wrapper.append(parent_wrap)

                            else:

                                if field_wrap_obj(filtered_obj[1], False) not in parent_wrapper[check_obj_list.index(
                                        parent_wrap.objName)].fieldWrapperList:
                                    parent_wrapper[check_obj_list.index(
                                        parent_wrap.objName
                                    )].fieldWrapperList.append(field_wrap_obj(filtered_obj[1], False))
                                    parent_wrapper[check_obj_list.index(
                                        parent_wrap.objName
                                    )].grandObjWrapperList.append(grand_wrap_obj(filtered_obj[1], False, [field_wrap_obj(filtered_obj[2], False)]))
                                else:

                                    check_grandobj_list = list()
                                    if 'grandObjWrapperList' in parent_wrapper[
                                            check_obj_list.index(
                                                parent_wrap.objName)]:

                                        for obj in parent_wrapper[check_obj_list.index(
                                                parent_wrap.objName
                                        )].grandObjWrapperList:
                                            check_grandobj_list.append(
                                                obj.objName)
                                    else:
                                        check_grandobj_list = []
                                    if grand_wrap.objName not in check_grandobj_list:
                                        grand_wrapper.append(grand_wrap)
                                    else:
                                        if field_wrap_obj(filtered_obj[2], False) not in parent_wrapper[check_obj_list.index(
                                                parent_wrap.objName
                                        )].grandObjWrapperList[
                                                check_grandobj_list.index(
                                                    grand_wrap.objName
                                                )].fieldWrapperList:

                                            parent_wrapper[check_obj_list.index(
                                                parent_wrap.objName
                                            )].grandObjWrapperList[
                                                check_grandobj_list.index(
                                                    grand_wrap.objName
                                                )].fieldWrapperList.append(field_wrap_obj(filtered_obj[2], False))

                        else:
                            parent_wrapper.append(parent_wrap)

                        grand_wrap = {}
                        grand_parent_field_wrap = {}
                        parent_field_wrap = {}
                        parent_wrap = {}
                        parent_field_wrapper = []
                        grand_parent_field_wrapper = []
                        grand_wrapper = []

                    elif len(filtered_obj) == 2:
                        parent_field_wrap = field_wrap_obj(
                            filtered_obj[-1], False)
                        check_parent_obj_list = list()
                        for obj in parent_wrapper:
                            check_parent_obj_list.append(obj.objName)
                        if ({
                                'objName': filtered_obj[0]
                        } in check_parent_obj_list
                        ) and parent_field_wrap not in parent_wrapper[
                                check_parent_obj_list.index(
                                    filtered_obj[0])]['fieldWrapperList']:
                            parent_field_wrapper.append(parent_field_wrap)

                        parent_wrap = parent_wrap_obj(
                            filtered_obj[0], False, parent_field_wrapper, [], [], [])
                        if len(parent_wrapper) > 0:
                            check_obj_list = list()
                            for obj in parent_wrapper:
                                check_obj_list.append(obj.objName)
                            if parent_wrap.objName not in check_obj_list:
                                parent_wrap.fieldWrapperList = [
                                    field_wrap_obj(filtered_obj[-1], False)]
                                parent_wrapper.append(parent_wrap)
                                parent_field_wrapper = []

                            elif {
                                    'fieldName': filtered_obj[-1],
                                    'isExists': False
                            } not in parent_wrapper[check_obj_list.index(
                                    parent_wrap.objName)].fieldWrapperList:
                                parent_wrapper[check_obj_list.index(
                                    parent_wrap.objName
                                )].fieldWrapperList.append(field_wrap_obj(filtered_obj[-1], False))
                                parent_field_wrapper = []
                        else:
                            parent_wrap.fieldWrapperList.append(
                                field_wrap_obj(filtered_obj[-1], False))
                            parent_wrapper.append(parent_wrap)
                            parent_field_wrapper = []
                    parent_field_wrapper = []

            obj_wrapper.fieldWrapperList = field_wrapper
            obj_wrapper.parentObjWrapperList = parent_wrapper

        parent_wrapper = []
        parent_field_wrapper = []
        parent_field_wrap = {}
        parent_wrap = {}
        old_child_obj_meta = []
        if len(child_obj_metadata) > 0:
            child_obj_wrapper_list = utility_class.form_child_obj_fields(
                child_obj_metadata, pick_list_fields, child_table_list)
            print("child_obj_wrapper_list-->{}".format(child_obj_wrapper_list))
            print("obj_wrapper_list-->{}".format(json.dumps(obj_wrapper,
                                                            default=lambda o: o.__dict__)))
            data = {
                "recordId": record_id,
                "jsonData": obj_wrapper,
                "folderId": folder_id_dyn,
                "fileName": file_name
            }
            obj_wrapper = json.dumps(obj_wrapper, default=lambda o: o.__dict__)
            json_dic = json.loads(obj_wrapper)
            object_name_new = json_dic.get('objName')
            collection_name = db[object_name_new.split(
                '#')[1]+'.'+object_name_new.split('#')[2]]
            queried_data = collection_name.aggregate(generate_mongo_query.form_mongo_query(
                json_dic, record_id, pick_list_fields, '_id'))
            mainobj_record_data = list(queried_data)[0]
            for child_wrapper in child_obj_wrapper_list:
                child_wrapper = json.dumps(
                    child_wrapper, default=lambda o: o.__dict__)
                json_dic = json.loads(child_wrapper)
                object_name_new = json_dic.get('objName')
                collection_name = db[object_name_new.split(
                    '#')[1]+'.'+object_name_new.split('#')[2]]
                get_obj_name = ''
                for pattern in child_table_pattern_list:
                    if re.search(object_name_new.split('#')[2], pattern):
                        get_obj_name = pattern
                replace_name = get_obj_name.split('(')
                queried_data = collection_name.aggregate(generate_mongo_query.form_mongo_query(
                    json_dic, record_id, pick_list_fields, replace_name[1].replace(')', '')))
                child_record_data = list(queried_data)
                mainobj_record_data[object_name_new] = {
                    'records': child_record_data}
            print('mainobj_record_data-->', mainobj_record_data)
            bind_values_doc(mainobj_record_data,
                            doc, table_pattern_list)
            # Testing purpose
            doc.save('test.docx')
            docx_stream = io.BytesIO()
            doc.save(docx_stream)
            docx_bytes = docx_stream.getvalue()
            encoded = base64.b64encode(docx_bytes)
            doc_data = {"isSuccess": True, "data": {
                "file_id": file_id,
                "file": str(encoded)[2:-1],
                "record_id": record_id,
                "success": True
            }}
            return json.dumps(doc_data)
        else:
            obj_wrapper = json.dumps(obj_wrapper, default=lambda o: o.__dict__)
            print("ObjMetaDataInfo-->{}".format(obj_wrapper))
            data = {
                "recordId": record_id,
                "jsonData": obj_wrapper,
                "folderId": folder_id_dyn,
                "fileName": file_name
            }
            json_dic = json.loads(obj_wrapper)
            object_name_new = json_dic.get('objName')
            collection_name = db[object_name_new.split(
                '#')[1]+'.'+object_name_new.split('#')[2]]
            queried_data = collection_name.aggregate(generate_mongo_query.form_mongo_query(
                json_dic, record_id, pick_list_fields, '_id'))
            record_data = list(queried_data)[0]
            for obj_field in record_data:
                if type(record_data[obj_field]) in (tuple, list):
                    new_value = {'records': record_data[obj_field]}
                    record_data[obj_field] = new_value
            print('queried_data-->', collection_name.aggregate(generate_mongo_query.form_mongo_query(
                json_dic, record_id, pick_list_fields, '_id')))

            # Bind values to the document
            bind_values_doc(record_data, doc, table_pattern_list)

            # Testing purpose
            doc.save('test.docx')
            docx_stream = io.BytesIO()
            doc.save(docx_stream)

            docx_bytes = docx_stream.getvalue()
            encoded = base64.b64encode(docx_bytes)
            doc_data = {"isSuccess": True, "data": {
                "file_id": file_id,
                "file": str(encoded)[2:-1],
                "record_id": record_id,
                "success": True
            }}
            return json.dumps(doc_data)
    except Exception as e:
        error_obj = {"isSuccess": False, "message":  str(e)}
        print(traceback.format_exc())
        return json.dumps(error_obj)


# To bind values to the fields which are not inside the table in the document
def bind_values_doc(data_dict, doc, table_pattern_list):
    utility_class = utility.utility()
    full_text_after = []
    child_obj_metadata = []
    for para in doc.paragraphs:
        full_text_after.append(para.text)
    document_data_after = '\n'.join(full_text_after)

    withouttable = utility_class.get_all_table_patterns(
        document_data_after.replace('\n', ' ').replace('\r', ''), table_pattern_list)
    if len(withouttable) > 0:
        child_obj_metadata = re.findall("\\$\\{(.*?)\\}", withouttable[0])
    for paragraph in doc.paragraphs:
        adjust_pattern = re.findall("\\{{ADJUST\\:(.*?)\\}}", paragraph.text)
        if len(adjust_pattern) > 0:
            matched_patterns = re.findall("\\$\\{(.*?)\\}", adjust_pattern[0])
            format_type = re.findall("\\((.*?)\\)", adjust_pattern[0])[0]
            format_type = format_type.split(',')
            date_value = utility_class.attach_field_values(
                matched_patterns[0], data_dict)
            separate_date = date_value.split('-')
            datefield = separate_date[2][:2]
            value = date(int(separate_date[0]), int(
                separate_date[1]), int(datefield))
            value = value + datetime.timedelta(int(format_type[1])*365/12)
            value = value + datetime.timedelta(days=int(format_type[0]))
            value = value + datetime.timedelta(int(format_type[2])*365)
            text_in_cell = paragraph.text
            value = text_in_cell.replace(
                '{{ADJUST:'+adjust_pattern[0]+'}}', str(value))
            paragraph.text = value
            target_stream = StringIO()
        matched_patterns = re.findall("\\$\\{(.*?)\\}", paragraph.text)
        function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", paragraph.text)
        if len(matched_patterns) > 0 and matched_patterns[0] not in child_obj_metadata:
            for value in matched_patterns:
                text_in_cell = paragraph.text
                field_value = utility_class.attach_field_values(
                    value, data_dict)
                field_value = text_in_cell.replace(
                    '${'+value+'}', str(field_value))
                paragraph.text = field_value
                target_stream = StringIO()
        if len(function_list) > 0:
            field_value = ''
            field_value = utility_class.generate_functions(
                function_list, data_dict)
            paragraph.text = field_value
        target_stream = StringIO()
        # doc.save(doc)

    para_table_obj = []
    child_tbl_objs = set()
    # Bind values outside child table
    for paragraph in doc.paragraphs:
        child_obj_metadata = re.findall("\\$\\{(.*?)\\}", paragraph.text)
        if len(child_obj_metadata) > 0:
            child_obj_fields = child_obj_metadata[0].split('.')
            child_tbl_objs.add(child_obj_fields[1])

    # Bind SUM values outside the cild table
    for cell in doc.paragraphs:
        count_func_list = re.findall("\\{\\{RowCount:(.*?)\\}}", cell.text)
        has_sum_func = re.findall("\\SUM\\{(.*?)\\}", cell.text)
        if len(count_func_list) > 0:
            for value in count_func_list:
                text_in_cell = cell.text
                field_value = str(
                    len(data_dict[count_func_list[0]]['records']))
                field_value = text_in_cell.replace(
                    '{{RowCount:'+value+'}}', field_value)
                cell.text = field_value
        elif len(has_sum_func) > 0:
            format_type = re.findall("(#[A-Z]*)", has_sum_func[0])
            form_org = has_sum_func[0].split('#')
            split_objname = form_org[2].split('.')
            print('form_org-->', form_org)
            print('split_objname-->', split_objname)
            form_childname = '#'+form_org[1]+'#' + \
                split_objname[1]+'('+split_objname[0]+')'
            splited_list = has_sum_func[0].split('.')
            if len(format_type) > 0:
                corrected_field = ''
                if len(format_type) > 0:
                    corrected_field = splited_list[-1].replace(
                        format_type[0], '').rstrip()
                else:
                    corrected_field = splited_list[-1]
                sum_of_field = 0
                formatted_type_data = ''
                for field in data_dict[form_childname]['records']:
                    if len(splited_list) > 0:
                        if len(splited_list) == 3:
                            try:
                                formatted_type_data = str(
                                    field[corrected_field])
                            except:
                                formatted_type_data = ''
                            formatted_type = str(formatted_type_data)
                        elif len(splited_list) == 4:
                            obj_name_match = re.split('Id', splited_list[2])
                            try:
                                formatted_type_data = field[obj_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]].keys(
                                ) else ''
                            except:
                                formatted_type_data = ''
                            formatted_type = str(formatted_type_data)
                        elif len(splited_list) == 5:
                            obj_name_match = re.split('Id', splited_list[2])
                            field_name_match = re.split('Id', splited_list[3])
                            try:
                                formatted_type_data = field[obj_name_match[0]][field_name_match[0]
                                                                               ][corrected_field] if corrected_field in field[obj_name_match[0]][field_name_match].keys() else ''
                            except:
                                formatted_type_data = ''
                            formatted_type = str(formatted_type_data)
                    sum_of_field = sum_of_field + float(formatted_type_data)
                    # sum_of_field = sum_of_field + float(field[corrected_field])
                text_in_cell = cell.text
                curr_value = locale.currency(sum_of_field, grouping=True)
                value = curr_value
                field_value = text_in_cell.replace(
                    '$SUM{'+has_sum_func[0]+'}', value)
                cell.text = field_value
            else:
                corrected_field = ''
                if len(format_type) > 0:
                    corrected_field = splited_list[-1].replace(
                        format_type[0], '').rstrip()
                else:
                    corrected_field = splited_list[-1]
                sum_of_field = 0
                formatted_type_data = ''
                for field in data_dict[splited_list[1]]['records']:
                    if len(splited_list) > 0:
                        if len(splited_list) == 3:
                            formatted_type = str(field[corrected_field])
                        elif len(splited_list) == 4:
                            obj_name_match = re.split('Id', splited_list[2])
                            formatted_type_data = field[obj_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]].keys(
                            ) else ''
                            formatted_type = str(formatted_type_data)
                        elif len(splited_list) == 5:
                            obj_name_match = re.split('Id', splited_list[2])
                            field_name_match = re.split('Id', splited_list[3])
                            try:
                                formatted_type_data = field[obj_name_match[0]][field_name_match][
                                    corrected_field] if corrected_field in field[obj_name_match[0]][field_name_match].keys() else ''
                            except:
                                formatted_type_data = ''
                            formatted_type = str(formatted_type_data)
                    sum_of_field = sum_of_field + float(formatted_type)
                text_in_cell = cell.text
                field_value = str(sum_of_field)
                field_value = text_in_cell.replace(
                    '$SUM{'+has_sum_func[0]+'}', field_value)
                cell.text = field_value
    table_obj_to_bind_list = []
    for objects in child_tbl_objs:
        fields_list = []
        for paragraph in doc.paragraphs:
            child_obj_metadata = re.findall("\\$\\{(.*?)\\}", paragraph.text)
            if len(child_obj_metadata) > 0:
                child_obj_fields = child_obj_metadata[0].split('.')
                if objects == child_obj_fields[1]:
                    fields_list.append(child_obj_metadata[0])
        table_obj_to_bind_list.append(
            {'objName': objects, 'fieldList': fields_list})

    for just_iterate in table_obj_to_bind_list:
        for paragraph in doc.paragraphs:
            table_obj = re.findall("\\$tbl{START:(.*):", paragraph.text)
            if len(table_obj) == 0:
                table_obj = re.findall(
                    "\\$tbl\\{START:(.*?)\\}", paragraph.text)
            if len(table_obj) > 0:
                if just_iterate['objName'] == table_obj[0]:
                    for record in data_dict[table_obj[0]]['records']:
                        for fields in just_iterate['fieldList']:
                            child_obj_fields = fields.split('.')
                            field_pattern = fields.split('.')
                            format_type = re.findall("(#[A-Z]*)", fields)
                            corrected_field = ''
                            if len(format_type) > 0:
                                corrected_field = field_pattern[-1].replace(
                                    format_type[0], '').rstrip()
                            else:
                                corrected_field = field_pattern[-1]
                                splited_list = corrected_field.split('.')
                            if len(field_pattern) == 3:
                                try:
                                    formatted_type_data = record[corrected_field]
                                except:
                                    formatted_type_data = ''
                                formatted_type = str(formatted_type_data)
                            elif len(field_pattern) == 4:
                                obj_name_match = re.split(
                                    'Id', field_pattern[2])
                                try:
                                    formatted_type_data = record[obj_name_match[0]][corrected_field] if corrected_field in record[obj_name_match[0]].keys(
                                    ) else ''
                                except:
                                    formatted_type_data = ''
                                formatted_type = str(formatted_type_data)
                            elif len(field_pattern) == 5:
                                obj_name_match = re.split(
                                    'Id', field_pattern[2])
                                field_name_match = re.split(
                                    'Id', field_pattern[3])
                                try:
                                    formatted_type_data = record[obj_name_match[0]][field_name_match[0]
                                                                                    ][corrected_field] if corrected_field in record[obj_name_match[0]][field_name_match[0]].keys() else ''
                                except:
                                    formatted_type_data = ''
                                formatted_type = str(formatted_type_data)
                            if len(format_type) > 0 and format_type[0] == ' #NUMBER':
                                value = ','.join(
                                    formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                            elif len(format_type) > 0 and format_type[0] == ' #CURRENCY':
                                curr_value = locale.currency(
                                    formatted_type_data, grouping=True)
                                value = curr_value
                            elif len(format_type) > 0 and format_type[0] == ' #DATE':
                                separate_date = formatted_type.split('-')
                                datefield = separate_date[2][:2]
                                value = date(int(separate_date[0]), int(
                                    separate_date[1]), int(datefield)).ctime()
                                value = value.split(' ')
                                value = value[1]+' '+value[2]+','+''+value[-1]
                            field_name = value if len(
                                format_type) > 0 else formatted_type
                            paragraph.insert_paragraph_before(field_name)
                            target_stream = StringIO()

    for just_iterate in table_obj_to_bind_list:
        for paragraph in doc.paragraphs:
            table_obj = re.findall("\\$tbl{START:(.*):", paragraph.text)
            table_end = re.findall("\\$tbl\\{END:(.*?)\\}", paragraph.text)
            if len(table_obj) == 0:
                table_obj = re.findall(
                    "\\$tbl\\{START:(.*?)\\}", paragraph.text)
            if len(table_obj) > 0:
                if just_iterate['objName'] == table_obj[0]:
                    paragraph.text = ''
            elif len(table_end) > 0:
                if just_iterate['objName'] == table_end[0]:
                    paragraph.text = ''
            else:
                child_obj_metadata = re.findall(
                    "\\$\\{(.*?)\\}", paragraph.text)
                if len(child_obj_metadata) > 0:
                    child_obj_fields = child_obj_metadata[0].split('.')
                    if just_iterate['objName'] == child_obj_fields[1]:
                        paragraph.text = ''
    target_stream = StringIO()

    alltext_in_tbl = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                alltext_in_tbl.append(cell.text)
    alltext_in_tbl = '\n'.join(alltext_in_tbl)
    table_patterns = re.findall(
        "\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:.*?\\}", alltext_in_tbl.replace('\n', ' ').replace('\r', ''))
    table_pattern_list = utility_class.get_all_table_patterns(
        alltext_in_tbl.replace('\n', ' ').replace('\r', ''), table_pattern_list)
    table_pattern_string = ' '.join(table_pattern_list)
    print('table_pattern_string-->', table_pattern_string)
    print('table_patterns-->', table_patterns)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matched_patterns = re.findall("\\$\\{(.*?)\\}", cell.text)
                function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", cell.text)
                field_value = ''
                if len(function_list) > 0:
                    field_value = utility_class.generate_functions(
                        function_list, data_dict)
                    cell.text = field_value
                elif len(matched_patterns) > 0:
                    for value in matched_patterns:

                        if len(table_patterns) > 0 and matched_patterns[0] in table_pattern_string:
                            # print('matched_patterns-->', matched_patterns)
                            # text_in_cell = cell.text
                            # field_value = utility_class.attach_field_values(
                            #     value, data_dict)
                            # print('field_value-->', field_value)
                            # field_value = text_in_cell.replace(
                            #     '${'+value+'}', str(field_value))
                            # cell.text = field_value
                            pass
                        else:
                            text_in_cell = cell.text
                            field_value = utility_class.attach_field_values(
                                value, data_dict)
                            field_value = text_in_cell.replace(
                                '${'+value+'}', str(field_value))
                            cell.text = field_value

    target_stream = StringIO()

    # Iterating tables to bind parent field values
    if len(doc.tables) > 0:
        table_fields_list = []
        alltext_in_doc = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    alltext_in_doc.append(cell.text)
                    # cell.text = attach_field_values(cell.text,data_dict,file_path)
                    matched_patterns = re.findall("\\$\\{(.*?)\\}", cell.text)
                    function_list = re.findall(
                        "\\{\\{FUNC:(.*?)\\}}", cell.text)

                    if len(matched_patterns) > 0:
                        if len(table_patterns) > 0 and matched_patterns[0] in table_pattern_string:
                            pass
                        else:
                            for value in matched_patterns:
                                cell.text = utility_class.attach_field_values(
                                    value, data_dict)
                    elif len(function_list) > 0:
                        field_value = ''
                        field_value = utility_class.generate_functions(
                            function_list, data_dict)
                        cell.text = field_value
                    target_stream = StringIO()

                    # doc.save(doc)
        alltext_in_doc = '\n'.join(alltext_in_doc)
        table_values = re.findall(
            "\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:.*?\\}", alltext_in_doc.replace('\n', ' ').replace('\r', ''))

        if len(table_values) > 0:
            table_fields_list = re.findall(
                "\\$\\{(.*?)\\}", table_pattern_string)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        table_row_details = re.findall(
                            "\\$\\{(.*?)\\}", cell.text)
                        count_func_list = re.findall(
                            "\\{\\{RowCount:(.*?)\\}}", cell.text)
                        has_sum_func = re.findall(
                            "\\SUM\\{(.*?)\\}", cell.text)
                        if len(count_func_list) > 0:
                            for value in count_func_list:
                                text_in_cell = cell.text
                                field_value = str(
                                    len(data_dict[count_func_list[0]]['records']))
                                field_value = text_in_cell.replace(
                                    '{{RowCount:'+value+'}}', field_value)
                                cell.text = field_value
                        elif len(has_sum_func) > 0:
                            format_type = re.findall(
                                "(#[A-Z]*)", has_sum_func[0])
                            splited_list = has_sum_func[0].split('.')
                            if len(format_type) > 0:
                                corrected_field = ''
                                if len(format_type) > 0:
                                    corrected_field = splited_list[-1].replace(
                                        format_type[0], '').rstrip()
                                else:
                                    corrected_field = splited_list[-1]
                                sum_of_field = 0
                                formatted_type_data = ''
                                for field in data_dict[splited_list[1]]['records']:
                                    if len(splited_list) > 0:
                                        if len(splited_list) == 3:
                                            formatted_type = str(
                                                field[corrected_field])
                                            formatted_type_data = field[corrected_field]
                                        elif len(splited_list) == 4:
                                            obj_name_match = re.split(
                                                'Id', splited_list[2])
                                            formatted_type_data = field[obj_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]].keys(
                                            ) else ''
                                            formatted_type = str(
                                                formatted_type_data)
                                        elif len(splited_list) == 5:
                                            obj_name_match = re.split(
                                                'Id', splited_list[2])
                                            field_name_match = re.split(
                                                'Id', splited_list[3])
                                            try:
                                                formatted_type_data = field[obj_name_match[0]][field_name_match[0]][
                                                    corrected_field] if corrected_field in field[obj_name_match[0]][field_name_match].keys() else ''
                                            except:
                                                formatted_type_data = ''
                                            formatted_type = str(
                                                formatted_type_data)
                                    sum_of_field = sum_of_field + \
                                        float(formatted_type_data)
                                    # sum_of_field = sum_of_field + float(field[corrected_field])
                                text_in_cell = cell.text
                                curr_value = locale.currency(
                                    sum_of_field, grouping=True)
                                value = curr_value
                                field_value = text_in_cell.replace(
                                    '$SUM{'+has_sum_func[0]+'}', value)
                                cell.text = field_value
                            else:
                                corrected_field = ''
                                if len(format_type) > 0:
                                    corrected_field = splited_list[-1].replace(
                                        format_type[0], '').rstrip()
                                else:
                                    corrected_field = splited_list[-1]
                                sum_of_field = 0
                                formatted_type_data = ''
                                for field in data_dict[splited_list[1]]['records']:
                                    if len(splited_list) > 0:
                                        if len(splited_list) == 3:
                                            try:
                                                formatted_type_data = str(
                                                    field[corrected_field])
                                            except:
                                                formatted_type_data = ''
                                            formatted_type = str(
                                                formatted_type_data)
                                        elif len(splited_list) == 4:
                                            obj_name_match = re.split(
                                                'Id', splited_list[2])
                                            try:
                                                formatted_type_data = field[obj_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]].keys(
                                                ) else ''
                                            except:
                                                formatted_type_data = ''
                                            formatted_type = str(
                                                formatted_type_data)
                                        elif len(splited_list) == 5:
                                            obj_name_match = re.split(
                                                'Id', splited_list[2])
                                            field_name_match = re.split(
                                                'Id', splited_list[3])
                                            try:
                                                formatted_type_data = field[obj_name_match[0]][field_name_match][
                                                    corrected_field] if corrected_field in field[obj_name_match[0]][field_name_match].keys() else ''
                                            except:
                                                formatted_type_data = ''
                                            formatted_type = str(
                                                formatted_type_data)
                                    sum_of_field = sum_of_field + \
                                        float(formatted_type)
                                text_in_cell = cell.text
                                field_value = str(sum_of_field)
                                field_value = text_in_cell.replace(
                                    '$SUM{'+has_sum_func[0]+'}', field_value)
                                cell.text = field_value

                        if len(table_row_details) > 0 and table_row_details[0] not in table_fields_list:
                            matched_patterns = re.findall(
                                "\\$\\{(.*?)\\}", cell.text)
                            function_list = re.findall(
                                "\\{\\{FUNC:(.*?)\\}}", cell.text)
                            field_value = ''
                            if len(function_list) > 0:
                                field_value = utility_class.generate_functions(
                                    function_list, data_dict)
                                cell.text = field_value
                            elif len(matched_patterns) > 0:
                                table_obj = re.findall(
                                    "\\$tbl\\{START:(.*?)\\}", cell.text)
                                for value in matched_patterns:
                                    text_in_cell = cell.text
                                    field_value = utility_class.attach_field_values(
                                        value, data_dict)
                                    field_value = text_in_cell.replace(
                                        '${'+value+'}', field_value)
                                    cell.text = field_value
            target_stream = StringIO()

        def remove_row(table, row):
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)

        # Iterating tables to bind child field values
        for table in doc.tables:
            column_value_list = []
            head_obj = []
            row_to_add = []
            for row_index, row in enumerate(table.rows):
                for column_index, cell in enumerate(row.cells):
                    check_child = re.findall("\\$\\{(.*?)\\}", cell.text)
                    if len(check_child) > 0 and check_child[0] in table_fields_list:
                        table_row_details = re.findall(
                            "\\$\\{(.*?)\\}", cell.text)
                    table_obj = re.findall("\\$tbl{START:(.*):", cell.text)
                    if len(table_obj) == 0:
                        table_obj = re.findall(
                            "\\$tbl\\{START:(.*?)\\}", cell.text)
                    table_end = re.findall("\\$tbl\\{END:(.*?)\\}", cell.text)
                    if len(table_obj) > 0:
                        row_to_add = table.row_cells(row_index)
                        head_obj = re.findall("\\$tbl{START:(.*):", cell.text)
                        if len(head_obj) == 0:
                            head_obj = re.findall(
                                "\\$tbl\\{START:(.*?)\\}", cell.text)
                        head_obj[0] = head_obj[0].strip()
                        row_columns = []
                        for cell in row_to_add:
                            if cell.text not in row_columns:
                                row_columns.append(cell.text)
                        if len(head_obj) > 0:
                            obj_name = head_obj[0].split('.')
                            check_is_child_obj = '#' + \
                                obj_name[0].split(
                                    '#')[1]+'#'+obj_name[1].split('(')[0]
                            if len(check_child) > 0 and check_child[0] in table_fields_list:
                                for i, record in enumerate(data_dict[check_is_child_obj]['records']):
                                    current_row = table.rows[row_index]
                                    border_copied = copy.deepcopy(
                                        current_row._tr)
                                    tr = border_copied
                                    current_row._tr.addnext(tr)
                                    for j, column in enumerate(row_columns):
                                        table_pattern = re.findall(
                                            "\\$\\{(.*?)\\}", column)
                                        if len(table_pattern) > 0:

                                            table_pattern = table_pattern[0].split(
                                                ':')
                                            field_pattern = table_pattern[1].split(
                                                '.')
                                            format_type = re.findall(
                                                "( #[A-Z]*)", table_pattern[1])
                                            corrected_field = ''
                                            if len(format_type) > 0:
                                                corrected_field = field_pattern[-1].replace(
                                                    format_type[0].strip(), '').strip()
                                            else:
                                                corrected_field = field_pattern[-1]
                                                splited_list = corrected_field.split(
                                                    '.')
                                            if len(field_pattern) == 2:
                                                try:
                                                    formatted_type_data = record[corrected_field]
                                                except:
                                                    formatted_type_data = ''
                                                formatted_type = str(
                                                    formatted_type_data)
                                            if len(field_pattern) == 3:
                                                try:
                                                    formatted_type_data = record[field_pattern[1]
                                                                                 ][corrected_field]
                                                except:
                                                    formatted_type_data = ''
                                                formatted_type = str(
                                                    formatted_type_data)
                                            elif len(field_pattern) == 4:
                                                obj_name_match = re.split(
                                                    'Id', field_pattern[2])
                                                try:
                                                    formatted_type_data = record[obj_name_match[0]][corrected_field] if corrected_field in record[obj_name_match[0]].keys(
                                                    ) else ''
                                                except:
                                                    formatted_type_data = ''
                                                formatted_type = str(
                                                    formatted_type_data)
                                            elif len(field_pattern) == 5:
                                                obj_name_match = re.split(
                                                    'Id', field_pattern[2])
                                                field_name_match = re.split(
                                                    'Id', field_pattern[3])
                                                try:
                                                    formatted_type_data = record[obj_name_match[0]][field_name_match[0]][
                                                        corrected_field] if corrected_field in record[obj_name_match[0]][field_name_match[0]].keys() else ''
                                                except:
                                                    formatted_type_data = ''
                                                formatted_type = str(
                                                    formatted_type_data)
                                            if len(format_type) > 0 and format_type[0] == ' #NUMBER':
                                                value = ','.join(
                                                    formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                                            elif len(format_type) > 0 and format_type[0] == ' #CURRENCY':
                                                curr_value = locale.currency(
                                                    formatted_type_data, grouping=True)
                                                value = curr_value
                                            elif len(format_type) > 0 and format_type[0] == ' #DATE':
                                                separate_date = formatted_type.split(
                                                    '-')
                                                datefield = separate_date[2][:2]
                                                value = date(int(separate_date[0]), int(
                                                    separate_date[1]), int(datefield)).ctime()
                                                value = value.split(' ')
                                                value = value[1]+' ' + \
                                                    value[2]+','+''+value[-1]
                                            field_name = value if len(
                                                format_type) > 0 else formatted_type
                                            table.cell(
                                                row_index+1, j).text = str(formatted_type_data)
                                            table.rows[row_index+1].height = 1
                    if len(table_end) > 0:
                        remove_row(table, table.rows[row_index])
        target_stream = StringIO()
        # doc.save(doc)


if __name__ == "__main__":
    app.run(host='0.0.0.0', port=port_number)

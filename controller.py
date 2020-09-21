import os
from flask import Flask, request, render_template, jsonify, url_for, redirect
from oauthlib.oauth2 import WebApplicationClient
import requests 
from werkzeug.utils import secure_filename 
from xml.etree import ElementTree
from docx import Document
from docx.document import Document as _Document
import docx
from base64 import b64decode
import re
import json  
from datetime import date
import datetime
import copy
import xml.etree.ElementTree as ET
from docx.enum.table import WD_ROW_HEIGHT
from Models.objwrapper import obj_wrap
from Models.fieldwrap import field_wrap_obj
from Models.grandwrap import grand_wrap_obj
from Models.parentwrap import parent_wrap_obj
from Models.childwrap import child_wrap_obj
from Models.user import user_credentials
from _io import BytesIO, StringIO
import base64
import io
import random
import string
import locale
locale.setlocale(locale.LC_ALL, 'en_US.utf-8')
from sqlalchemy import create_engine,Column,Integer,String,Date,ForeignKey,MetaData,Table
from sqlalchemy.ext.declarative import declarative_base 
import pymysql.cursors
from _datetime import date, timedelta
import uuid
import shutil
from flask_cors import CORS
import jwt
from oauthlib.common import urlencode


# Configuration
GOOGLE_CLIENT_ID = '956259232368-p6nm8n4ettu38anu3c9i5gpfh4610t0g.apps.googleusercontent.com'
GOOGLE_CLIENT_SECRET = '9eUivwpkG1__yz3MnTr016fG'
GOOGLE_DISCOVERY_URL = (
    "https://accounts.google.com/.well-known/openid-configuration"
)
client = WebApplicationClient(GOOGLE_CLIENT_ID)

# connectionObject   = pymysql.connect(host='localhost', user='root', password='Aspi@2018',db='python-doc-gen', charset='utf8mb4',cursorclass=pymysql.cursors.DictCursor)
# cursorObject        = connectionObject.cursor() 
# cursorObject.execute('CREATE TABLE IF NOT EXISTS `user` (`id` int NOT NULL AUTO_INCREMENT PRIMARY KEY,`username` varchar(45) DEFAULT NULL,`password` varchar(45) DEFAULT NULL,`organisationid` varchar(45) DEFAULT NULL,`email` varchar(45) DEFAULT NULL,`phonenumber` varchar(45) DEFAULT NULL,`accountid` varchar(45) DEFAULT NULL)')   
# cursorObject.execute('CREATE TABLE IF NOT EXISTS `userlog` (`id` int NOT NULL AUTO_INCREMENT PRIMARY KEY,`username` varchar(45) DEFAULT NULL,`userid` varchar(45) DEFAULT NULL,`organisationid` varchar(45) DEFAULT NULL,`filename` varchar(45) DEFAULT NULL,`generateddate` Date DEFAULT NULL)')  
# Base = declarative_base()
# SQLALCHEMY_DATABASE_URI = 'mysql+pymysql://root:Aspi@2018@localhost:3306/python-doc-gen'
# DB_Engine = create_engine(SQLALCHEMY_DATABASE_URI)
# connection = DB_Engine.connect()


# class UserLog(Base):
#     __tablename__ = "UserLog"

#     id= Column('user_id',Integer, primary_key= True)
#     username = Column('user_name',String(122), unique= True)
#     organizationid = Column('organization_id',String(122), unique= True)
#     folderid = Column('folder_id',String(122), unique= True)
#     generateddate = Column('generated_date',Date,unique= False )
#     filename = Column('file_name',String(122),unique= False)

# class User(Base):
#     __tablename__ = "User"

#     id= Column('id',Integer, primary_key= True)
#     username = Column('username',String(122), unique= True)
#     email = Column('email',String(122), unique= True)
#     organizationid = Column('organizationid',String(122), unique= True)
#     phonenumber = Column('phonenumber',Date,unique= False )
#     password = Column('password',String(122),unique= False)

# meta = MetaData(DB_Engine)
 
# user_log = Table(
#    'UserLog', meta, 
#    Column('user_id', Integer, primary_key = True), 
#    Column('user_name', String), 
#    Column('organization_id', String), 
#     Column('folder_id', String),
#      Column('generated_date', Date),
#      Column('file_name', String),
# )

# user = Table(
#    'User', meta, 
#    Column('id', Integer, primary_key = True), 
#    Column('username', String), 
#    Column('email', String), 
#     Column('organizationid', String),
#      Column('phonenumber', Date),
#      Column('password', String),
# )


   
                                                            
UPLOAD_FOLDER = str(os.getcwd())+'/python-mailmerge/Document'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif'}
path = str(os.getcwd())+'/Document'
app = Flask(__name__)
cors = CORS(app, resources={"/*": {"origins": "*"}})
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER 
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
file_path = ''
document_data = ''
document_data_initial = ''
record_id = ''
file_name = ''
file_type = ''
table_pattern_list = []
folder_id_dyn = ''

# @app.route('/api/auth/signup', methods =['POST'])
# def save_user_credentials():
#     response_data = json.loads(request.data.decode('utf-8'))
#     user_details = user_credentials('',response_data['username'],response_data['email'],response_data['organisationid'],response_data['phonenumber'],response_data['password'],'')
#     id = uuid.uuid1().hex
#     user_details.accountid = id
#     sql = "INSERT INTO `user` (`id`, `username`, `email`, `organisationid`, `phonenumber`, `password`, `accountid`) VALUES (%s, %s, %s, %s, %s, %s, %s)"
#     cursorObject.execute(sql, (0,user_details.username,user_details.email,user_details.organisationid,user_details.phonenumber,user_details.password,user_details.accountid))
#     connectionObject.commit()
#     result = cursorObject.fetchall()
#     data = {
#             "message" : "Account created successfully!",
#         }
#     for i in result:
#         print(i)
#     return json.dumps(data)

# @app.route('/api/auth/sflogin', methods =['GET'])
# def login_with_sf_credentials():
#     paramsadded = {"redirect_uris": "http://localhost:4200/home","response_types":"code"}
#     r=requests.post("https://gautidomain-dev-ed.my.salesforce.com/services/oauth2/authorize",data=paramsadded, headers={"Content-Type": "application/json","Accept":"application/json","Authorization":"Bearer 6Cel800D6F000002Xu0g8886F000001yXrWeSLML4Y4hM6ENS9rY16SuCMNsTbJRSZsMNcfgWUu4nmj8sc26ZyfoBJHf4NMS2qt2VXZBZgD","Host":"gautidomain-dev-ed.my.salesforce.com"})
#     print(r.text)
#     print(r.links)
#     print(r.headers)
#     data = {
#             "message" : "Loging in to salesforce!"
#         }
#     return json.dumps(r)

# def get_google_provider_cfg():
#     return requests.get(GOOGLE_DISCOVERY_URL).json()

# @app.route("/api/auth/googlelogin", methods =['GET'])
# def login():
#     # Find out what URL to hit for Google login
#     google_provider_cfg = get_google_provider_cfg()
#     authorization_endpoint = google_provider_cfg["authorization_endpoint"]
#     print("Exits","https://127.0.0.1:5000/login" + "/callback")
#     # Use library to construct the request for Google login and provide
#     # scopes that let you retrieve user's profile from Google
#     print("client",client)
#     request_uri = client.prepare_request_uri(
#         authorization_endpoint,
#         redirect_uri="https://127.0.0.1:5000/login" + "/callback",
#         scope=["openid", "email", "profile"],
#     )
#     print("Exits",request_uri)
#     data = {
#             "url" : request_uri
#         }
#     return json.dumps(data)

# @app.route("/login/callback")
# def callback():
#     # Get authorization code Google sent back to you
#     code = request.args.get("code")
#     print("code-->",type(code))
#     print("code-->",code)
#     print("request.url-->",request.url)
#     print("request.base_url-->",request.base_url)
#     google_provider_cfg = get_google_provider_cfg()
#     token_endpoint = google_provider_cfg["token_endpoint"]
#     token_url, headers, body = client.prepare_token_request(
#     token_endpoint,
#     authorization_response=request.url,
#     redirect_url=request.base_url,
#     code=code
#     )
#     token_response = requests.post(
#         token_url,
#         headers=headers,
#         data=body,
#         auth=(GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET),
#     )
#     # Parse the tokens!
#     client.parse_request_body_response(json.dumps(token_response.json()))
#     userinfo_endpoint = google_provider_cfg["userinfo_endpoint"]
#     uri, headers, body = client.add_token(userinfo_endpoint)
#     userinfo_response = requests.get(uri, headers=headers, data=body)
#     print("userinfo_response",userinfo_response.json())
#     return redirect("https://localhost:4200/home")


# @app.route("/api/auth/googletoken", methods =['POST'])
# def getgoogleloginjwt():
#     # Get authorization code Google sent back to you
#     response_data = json.loads(request.data.decode('utf-8'))
#     codevalue = response_data['code']
#     google_provider_cfg = get_google_provider_cfg()
#     token_endpoint = google_provider_cfg["token_endpoint"]
#     token_url, headers, body = client.prepare_token_request(
#     token_endpoint,
#     redirect_url=response_data['redirect_uri'],
#     authorization_response=response_data['callbackurl'],
#     code=codevalue
#     )
#     token_response = requests.post(
#         token_url,
#         headers=headers,
#         data=body,
#         auth=(GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET),
#     )
#     client.parse_request_body_response(json.dumps(token_response.json()))
#     userinfo_endpoint = google_provider_cfg["userinfo_endpoint"]
#     uri, headers, body = client.add_token(userinfo_endpoint)
#     userinfo_response = requests.get(uri, headers=headers, data=body)
#     google_user_credentials = userinfo_response.json()
#     print("userinfo_response",userinfo_response.json())
#     cursorObject.execute("SELECT * FROM user WHERE username=%s", google_user_credentials['name'])
#     rows = cursorObject.fetchall()
#     print("rows-->",rows)
#     if len(rows) == 0 :
#         user_details = user_credentials('',google_user_credentials['name'],google_user_credentials['email'],'','','','')
#         id = uuid.uuid1().hex
#         user_details.accountid = id
#         sql = "INSERT INTO `user` (`id`, `username`, `email`, `organisationid`, `phonenumber`, `password`, `accountid`) VALUES (%s, %s, %s, %s, %s, %s, %s)"
#         cursorObject.execute(sql, (0,user_details.username,user_details.email,user_details.organisationid,user_details.phonenumber,user_details.password,user_details.accountid))
#         connectionObject.commit()
#         encoded_jwt = jwt.encode({'usernameOrEmail': google_user_credentials['name']}, 'secret', algorithm='HS256')
#         jwt_toke = encoded_jwt.decode("utf-8")
#         decoded_jwt = jwt.decode(encoded_jwt, 'secret', algorithms=['HS256'])
#         message = "SUCCESS"
#     else :
#         encoded_jwt = jwt.encode({'usernameOrEmail': google_user_credentials['name']}, 'secret', algorithm='HS256')
#         jwt_toke = encoded_jwt.decode("utf-8")
#         decoded_jwt = jwt.decode(encoded_jwt, 'secret', algorithms=['HS256'])
#         message = "SUCCESS"
#     data = {
#             "token" : jwt_toke,
#             "message" : message
#          }
#     return json.dumps(data)




# @app.route('/api/auth/signin', methods =['POST'])
# def user_login():
#     response_data = json.loads(request.data.decode('utf-8'))
#     print("response_data-->",response_data)
#     cursorObject.execute("SELECT * FROM user WHERE username=%s", response_data['usernameOrEmail'])
#     rows = cursorObject.fetchall()
#     if len(rows) > 0 and rows[0]['password'] == response_data['password'] :
#         encoded_jwt = jwt.encode({'usernameOrEmail': response_data['usernameOrEmail']}, 'secret', algorithm='HS256')
#         jwt_toke = encoded_jwt.decode("utf-8")
#         decoded_jwt = jwt.decode(encoded_jwt, 'secret', algorithms=['HS256'])
#         message = "SUCCESS"
#     else :
#         jwt_toke = ''
#         message = "ERROR"
#     data = {
#             "token" : jwt_toke,
#             "message" : message
#          }
#     return json.dumps(data)

# @app.route('/api/user/currentuser', methods =['POST'])
# def current_user():
#     encode_token = request.data.decode('utf-8')
#     decoded_jwt = jwt.decode(encode_token, 'secret', algorithms=['HS256'])
#     return json.dumps(decoded_jwt)

# @app.route('/api/user/userdetails', methods =['POST'])
# def get_user_details():
#     user_name = request.data.decode('utf-8')
#     cursorObject.execute("SELECT * FROM user WHERE username=%s", user_name)
#     rows = cursorObject.fetchall()
#     return json.dumps(rows[0])

# @app.route('/api/document/dashboardcount', methods =['POST'])
# def get_dashboard_count():
#     user_name = request.data.decode('utf-8')
#     cursorObject.execute("SELECT * FROM userlog WHERE username=%s", user_name)
#     total_records = cursorObject.fetchall()
#     last_week_date =  date.today() - timedelta(days=7)
#     last_month_date =  date.today() - timedelta(1*365/12)
#     cursorObject.execute("SELECT * FROM userlog WHERE username='" + user_name + "' AND  DATE(generateddate) >= " + str(last_week_date))
#     last_week_records = cursorObject.fetchall()
#     separate_date = str(last_month_date).split('-')
#     date_modified = datetime.date(int(separate_date[0]),int(separate_date[1]), int(separate_date[2]))
#     cursorObject.execute("SELECT * FROM userlog WHERE username='" + user_name + "' AND  CAST(generateddate AS DATE) >= " + str(date_modified))
#     last_month_records = cursorObject.fetchall()
#     print("last_month_date-->",date_modified)
#     print("last_month_date-->",last_month_records)
    
#     data = {
#             "totaldocuments" : len(total_records),
#             "lastweek" : len(last_week_records),
#             "lastmonth" : len(last_month_records)
#          }
#     return json.dumps(data)

@app.route('/')
def index():
    return render_template('indexlive.html')
    
@app.route('/handle_form', methods =['POST'])
def handle_form():
    content = request.files['file'].read()
    record_id = str(request.form.get('recordId'))
    file_name = request.files['file'].filename
    file_type = request.headers['fileType']
    file_name = file_name+'.docx'
    user_id = str(request.form.get('userId'))
    user_name = str(request.form.get('userName'))
    org_id = str(request.form.get('orgId'))
    # ins = user_log.insert().values(user_id = user_id, user_name = user_name,
    # organization_id = org_id,file_name = file_name,generated_date = date.today())
    # connection.execute(ins)
    # Base.metadata.create_all(bind=DB_Engine)
    bytes = b64decode(content)
    source_stream = BytesIO(content)
    doc = Document(source_stream)
    source_stream.close()
    full_text = []
    field_list = []
    child_obj_metadata = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    document_data = '\n'.join(full_text)
    field_list = re.findall("\\$\\{(.*?)\\}", document_data)
    child_obj_metadata = re.findall("\\$tbl\\{(.*?)\\}", document_data)
    withouttable = get_all_table_patterns(document_data.replace('\n', ' ').replace('\r', ''))
    if len(withouttable) > 0 : 
        child_obj_metadata = re.findall("\\$\\{(.*?)\\}", withouttable[0])
        field_list = list(set(field_list) - set(child_obj_metadata))
    head_child_obj = '' 
    alltext_in_tbl = []
    for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    alltext_in_tbl.append(cell.text)
    alltext_in_tbl = '\n'.join(alltext_in_tbl)
    table_patterns = re.findall("\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:[A-Za-z]*\\}", alltext_in_tbl.replace('\n', ' ').replace('\r', ''))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                fields_in_cell = re.findall("\\$\\{(.*?)\\}",cell.text)
                child_obj_values = re.findall("\\$tbl\\{(.*?)\\}",cell.text)
                child_obj_str = []
                if len(child_obj_values) > 0:
                   child_obj_str =  re.findall("\\$tbl\\{START:(.*?)\\}", cell.text) 
                if len(child_obj_str) > 0:
                   head_child_obj = child_obj_str[0]
                    
                # if len(fields_in_cell) > 0 and fields_in_cell[0].split('.')[1] == head_child_obj.strip():
                #     print("EXISTS")
                #     for field in fields_in_cell :
                #         child_obj_metadata.append(field)
                # elif len(fields_in_cell) > 0 :
                #     for field in fields_in_cell :
                #         field_list.append(field)
                if len(fields_in_cell) > 0 :
                    for field in fields_in_cell :
                        if field.split('.')[1] == head_child_obj.strip():
                            child_obj_metadata.append(field)
                        else:
                            field_list.append(field)
    field_list = list(dict.fromkeys(field_list))
    
    
    #This block will do json formation of fields and parent objects which will be returned to salesforce for retrieving data             
    if len(field_list) > 0 :
        obj_wrapper = obj_wrap(field_list[0].split('.')[0],False,[],[],[])
        field_wrapper = []
        parent_wrapper = []
        parent_field_wrapper = []
        grand_parent_field_wrapper = []
        grand_wrapper = []
        for field in field_list:
            format_type = re.findall("(#[A-Z]*)",field)
            if len(format_type) > 0 :
                corrected_field = field.replace(format_type[0].strip(),'')
            else :
                corrected_field = field
            parent_obj = corrected_field.split('.')
            if len(parent_obj) == 2:
                field_wrap = field_wrap_obj(parent_obj[-1],False)
                if field_wrap not in field_wrapper:
                    field_wrapper.append(field_wrap)

            else:
                field_wrap = field_wrap_obj(parent_obj[1],False)
                if field_wrap.__dict__ not in field_wrapper:
                    field_wrapper.append(field_wrap.__dict__)
                filtered_obj = parent_obj[1:len(parent_obj)]

                if len(filtered_obj) == 3:
                    parent_field_wrap = field_wrap_obj(filtered_obj[1],False)
                    if parent_field_wrap not in parent_field_wrapper:
                        parent_field_wrapper.append(parent_field_wrap)
                    grand_parent_field_wrap = field_wrap_obj(filtered_obj[2],False)
                    if grand_parent_field_wrap not in grand_parent_field_wrapper:
                        grand_parent_field_wrapper.append(grand_parent_field_wrap)
                    grand_wrap = grand_wrap_obj(filtered_obj[1],False,grand_parent_field_wrapper)
                    if len(grand_wrapper) > 0:
                        check_grobj_list = list()
                        for obj in grand_wrapper:
                            check_grobj_list.append(obj.objName)

                        if grand_wrap.objName not in check_grobj_list:
                            grand_wrapper.append(grand_wrap)

                        elif {
                                'fieldName': filtered_obj[2]
                        } not in grand_wrapper[check_grobj_list.index(
                                grand_wrap.objName)].fieldWrapperList:
                            grand_wrapper[check_grobj_list.index(
                                grand_wrap.objName)].fieldWrapperList.append(
                                    field_wrap_obj(filtered_obj[2],False))
                    else:
                        grand_wrapper.append(grand_wrap)
                    parent_wrap = parent_wrap_obj(filtered_obj[0],False,parent_field_wrapper,[],[],grand_wrapper)

                    if len(parent_wrapper) > 0:
                        check_obj_list = list()
                        for obj in parent_wrapper:
                            check_obj_list.append(obj.objName)
                        if parent_wrap.objName not in check_obj_list:
                            parent_wrapper.append(parent_wrap)

                        else:

                            if field_wrap_obj(filtered_obj[1],False) not in parent_wrapper[check_obj_list.index(
                                    parent_wrap.objName)].fieldWrapperList:
                                parent_wrapper[check_obj_list.index(
                                    parent_wrap.objName
                                )].fieldWrapperList.append(field_wrap_obj(filtered_obj[1],False))
                                # if 'grandObjWrapperList' in parent_wrapper[
                                #         check_obj_list.index(
                                #             filtered_obj[0])]:
                                parent_wrapper[check_obj_list.index(
                                    parent_wrap.objName
                                )].grandObjWrapperList.append(grand_wrap_obj(filtered_obj[1],False,[field_wrap_obj(filtered_obj[2],False)]))

                                # else:
                                # parent_wrapper[check_obj_list.index(
                                #     parent_wrap.objName
                                # )].grandObjWrapperList = [grand_wrap_obj(filtered_obj[1],False,[field_wrap_obj(filtered_obj[2],False)])]
                                
                                
                                
                                

                            else:

                                check_grandobj_list = list()
                                if 'grandObjWrapperList' in parent_wrapper[
                                        check_obj_list.index(
                                            parent_wrap.objName)]:

                                    for obj in parent_wrapper[check_obj_list.index(
                                            parent_wrap.objName
                                    )].grandObjWrapperList:
                                        check_grandobj_list.append(obj.objName)
                                else:
                                    check_grandobj_list = []
                                if grand_wrap.objName not in check_grandobj_list:
                                    grand_wrapper.append(grand_wrap)

                                else:

                                    if field_wrap_obj(filtered_obj[2],False) not in parent_wrapper[check_obj_list.index(
                                            parent_wrap.objName
                                    )].grandObjWrapperList[
                                            check_grandobj_list.index(
                                                grand_wrap.objName
                                            )].fieldWrapperList:

                                        parent_wrapper[check_obj_list.index(
                                            parent_wrap.objName
                                        )].grandObjWrapperList[
                                            check_grandobj_list.index(
                                                grand_wrap.objName
                                            )].fieldWrapperList.append(field_wrap_obj(filtered_obj[2],False))

                    else:
                        parent_wrapper.append(parent_wrap)

                    grand_wrap = {}
                    grand_parent_field_wrap = {}
                    parent_field_wrap = {}
                    parent_wrap = {}
                    parent_field_wrapper = []
                    grand_parent_field_wrapper = []
                    grand_wrapper = []

                elif len(filtered_obj) == 2:
                    parent_field_wrap = field_wrap_obj(filtered_obj[-1],False)
                    check_parent_obj_list = list()
                    for obj in parent_wrapper:
                        check_parent_obj_list.append(obj.objName)
                    if ({
                            'objName': filtered_obj[0]
                    } in check_parent_obj_list
                        ) and parent_field_wrap not in parent_wrapper[
                            check_parent_obj_list.index(
                                filtered_obj[0])]['fieldWrapperList']:
                        parent_field_wrapper.append(parent_field_wrap)

                    parent_wrap = parent_wrap_obj(filtered_obj[0],False,parent_field_wrapper,[],[],[])
                    if len(parent_wrapper) > 0:
                        check_obj_list = list()
                        for obj in parent_wrapper:
                            check_obj_list.append(obj.objName)
                        if parent_wrap.objName not in check_obj_list:
                            parent_wrap.fieldWrapperList = [field_wrap_obj(filtered_obj[-1],False)]
                            parent_wrapper.append(parent_wrap)
                            parent_field_wrapper = []

                        elif {
                                'fieldName': filtered_obj[-1],
                                'isExists': False
                        } not in parent_wrapper[check_obj_list.index(
                                parent_wrap.objName)].fieldWrapperList:
                            parent_wrapper[check_obj_list.index(
                                parent_wrap.objName
                            )].fieldWrapperList.append(field_wrap_obj(filtered_obj[-1],False))
                            parent_field_wrapper = []
                    else:
                        parent_wrap.fieldWrapperList.append(field_wrap_obj(filtered_obj[-1],False))
                        parent_wrapper.append(parent_wrap)
                        parent_field_wrapper = []
                parent_field_wrapper = []
        
        obj_wrapper.fieldWrapperList = field_wrapper
        obj_wrapper.parentObjWrapperList = parent_wrapper
        
        
       
        
        
    parent_wrapper = [] 
    parent_field_wrapper = []
    parent_field_wrap = {}
    parent_wrap = {}
    old_child_obj_meta = []
    
        
    
        
    
    #Method to return index of the obj 
    def check_obj_present(current_obj, child_wrapper):
        obj_list = list()
        for record in child_wrapper.parentObjWrapperList:
            obj_list.append(record.objName)
        if len(obj_list) > 0:
            try:
                return obj_list.index(current_obj)
            except ValueError:
                return -1
        else:
            return -1
    
    #Method to return index of the parent obj
    def check_grand_obj_present(current_obj, child_wrapper, parent_index):
        obj_list = list()

        for record in child_wrapper.parentObjWrapperList[parent_index].grandObjWrapperList:
            obj_list.append(record.objName)
        if len(obj_list) > 0:
            try:
                return obj_list.index(current_obj)
            except ValueError:
                return -1
        else:
            return -1
    #Method to return index of the field obj
    def check_field_obj_present(current_obj, child_wrapper):
        field_list = list()

        for record in child_wrapper.fieldWrapperList:
            field_list.append(record.fieldName)
        if len(field_list) > 0:
            try:
                return field_list.index(current_obj)
            except ValueError:
                return -1
        else:
            return -1
    
    #Method to generate childWrapper
    def generate_child_obj(child_obj, child_wrapper):
        child1_field_wrap = field_wrap_obj(child_obj[2],False)
        field_index = check_field_obj_present(child_obj[2], child_wrapper)
        if len(child_obj) == 3:
            if field_index == -1:
                    child_wrapper.fieldWrapperList.append(child1_field_wrap)
        if len(child_obj) == 4:
            if field_index == -1:
                child_wrapper.fieldWrapperList.append(child1_field_wrap)
            index_value = check_obj_present(child_obj[2], child_wrapper)
            if index_value == -1:
                child_wrapper.parentObjWrapperList.append(parent_wrap_obj(child_obj[2],False,[field_wrap_obj(child_obj[3],False)],[],[],[]))
            else:
                
                if field_wrap_obj(child_obj[3],False) not in child_wrapper.parentObjWrapperList[check_obj_present(child_obj[2],child_wrapper)].fieldWrapperList:
                    child_wrapper.parentObjWrapperList[check_obj_present(child_obj[2],child_wrapper)].fieldWrapperList.append(field_wrap_obj(child_obj[3],False))
        if len(child_obj) == 5:
            if field_index == -1:
                child_wrapper.fieldWrapperList.append(child1_field_wrap)
            index_value = check_obj_present(child_obj[2], child_wrapper)
            if index_value == -1:
                child_wrapper.parentObjWrapperList.append(parent_wrap_obj(child_obj[2],False,
                [field_wrap_obj(child_obj[3],False)],
                [grand_wrap_obj(child_obj[3],False,[field_wrap_obj(child_obj[4],False)])]
                ))
            else:
                if field_wrap_obj(child_obj[3],False) not in child_wrapper.parentObjWrapperList[index_value].fieldWrapperList:
                    child_wrapper.parentObjWrapperList[index_value].fieldWrapperList.append(field_wrap_obj(child_obj[3],False))
                grand_index_value = check_grand_obj_present(child_obj[3], child_wrapper, index_value)
                if grand_index_value == -1:
                    child_wrapper.parentObjWrapperList[index_value].grandObjWrapperList.append(grand_wrap_obj(child_obj[3],False,[field_wrap_obj(child_obj[4],False)]))
                else:
                    child_wrapper.parentObjWrapperList[index_value].grandObjWrapperList[grand_index_value].fieldWrapperList.append(field_wrap_obj(child_obj[4],False))
        return child_wrapper
    
    
            
    if len(child_obj_metadata) > 0:
        child_wrapper = child_wrap_obj('',False,[],[],'')
        child_obj_list = []
        child_wrapper_list = []
        get_condition = []
        check_whr_condition = re.findall("\\$tbl\\{END:(.*?)\\}",alltext_in_tbl)
        cnd_obj_set = dict()
        main_obj = ''
        if len(check_whr_condition) > 0 :
               for cnd_value in check_whr_condition :
                    try:
                        condition_on_table_idx = cnd_value.index('#CND')
                    except ValueError:
                        condition_on_table_idx = -1
                    if condition_on_table_idx != -1 :
                       cnd_table = cnd_value[condition_on_table_idx+4:]
                       cnd_obj = cnd_value[:condition_on_table_idx-1]
                       cnd_obj_set[cnd_obj]=cnd_table
                       
        for field in child_obj_metadata:
            format_type = re.findall("(#[A-Z]*)",field)
            if len(format_type) > 0 :
                corrected_field = field.replace(format_type[0],'')
            else :
                corrected_field = field
            child_obj = corrected_field.split('.')
            main_obj = child_obj[0]
            if child_obj not in old_child_obj_meta:
                if child_obj[1] not in child_obj_list:
                    head_obj = re.findall("\\$tbl{START:[A-Za-z]\\:(.*)",document_data)
                    if len(check_whr_condition) > 0 :
                        try :
                            child_obj_check = child_wrap_obj(child_obj[1],False,[],[],cnd_obj_set[child_obj[1]])
                        except KeyError:
                            child_obj_check = child_wrap_obj(child_obj[1],False,[],[],'')
                    else:
                        child_obj_check = child_wrap_obj(child_obj[1],False,[],[],'')
                    child_wrapper = generate_child_obj(child_obj,child_obj_check)
                    child_obj_list.append(child_obj[1])
                    child_wrapper_list.append(child_wrapper)
                else:
                    child_wrapper = generate_child_obj(child_obj,child_wrapper_list[child_obj_list.index(child_obj[1])])
                    child_wrapper_list[child_obj_list.index(child_obj[1])] = child_wrapper
                old_child_obj_meta.append(child_obj)
        if len(field_list) > 0 :
            obj_wrapper.childObjWrapperList = child_wrapper_list
        else:
            obj_wrapper = obj_wrap(main_obj,False,[],[],[])
            obj_wrapper.childObjWrapperList = child_wrapper_list
        obj_wrapper = json.dumps(obj_wrapper, default=lambda o: o.__dict__)
        print("ObjMetaDataInfo-->{}".format(obj_wrapper))
        data = {
            "recordId" : record_id,
            "jsonData" : obj_wrapper,
            "folderId": folder_id_dyn,
            "fileName" : file_name
        }
        return json.dumps(data)
    else :
        obj_wrapper = json.dumps(obj_wrapper, default=lambda o: o.__dict__)
        print("ObjMetaDataInfo-->{}".format(obj_wrapper))
        data = {
            "recordId" : record_id,
            "jsonData" : obj_wrapper,
            "folderId": folder_id_dyn,
            "fileName" : file_name
        }
        return json.dumps(data)


@app.route('/get_document', methods =['POST'])
def create_docx():
    content = request.files['file'].read()
    bytes = b64decode(content)
    source_stream = BytesIO(content)
    doc = Document(source_stream)
    source_stream.close()
    file_name = request.files['file'].filename
    record_id = str(request.form.get('recordId'))
    data_dict =  json.loads(request.form.get('recordData'))
    print("SalesforceData-->{}".format(data_dict))
    bind_values_doc(data_dict,doc)
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_bytes = docx_stream.getvalue()
    encoded = base64.b64encode(docx_bytes)
    doc_data = {
            "fileName":file_name ,
            "body": str(encoded)[2:-1],
            "parentId": record_id 
            }
    return json.dumps(doc_data)


#Create folder for document :
def create_folder(directory):
    try:
        folder_path = str(os.getcwd())
        os.chdir(folder_path)
        print("folder_path-->{}".format(folder_path))
        os.makedirs('Docu')
        print("Folder Created->")
        if not os.path.exists(directory):

            os.makedirs(directory)
    except OSError:
        print("Error: Creating directory. {}".format(directory))

#Mock response, instead of salesforce response
@app.route('/bind_document', methods=['POST'])
def get_data_sf() :
    r = request.data
    data_dict = json.loads(r.text)
    print("request.data-->{}".format(data_dict))
    return requests.get("http://www.mocky.io/v2/5e4631fc3300004d00025f7f")

def get_all_table_patterns(whole_text):
    table_patterns = re.findall("\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:.*?\\}",whole_text)
    if len(table_patterns) > 0:
        table_pattern_list.append(table_patterns[0])
        remaining_text = whole_text.index('END')
        print(remaining_text,whole_text[remaining_text+3:])
        if remaining_text != -1 :
            return get_all_table_patterns(whole_text[remaining_text+3:])
        else :
            return table_pattern_list
    else :
        return table_pattern_list

    
    

#To bind values to the fields which are not inside the table in the document
def bind_values_doc(data_dict,doc):
    # doc = docx.Document(file_path)
    # docume = Document(file_path)
    full_text_after = []
    child_obj_metadata = []
    for para in doc.paragraphs:
        full_text_after.append(para.text)
    document_data_after = '\n'.join(full_text_after)
    print('document_data_after-->',document_data_after)
    withouttable = get_all_table_patterns(document_data_after.replace('\n', ' ').replace('\r', ''))
    print('get_all_table_patterns-->',withouttable)
    if len(withouttable) > 0 : 
        child_obj_metadata = re.findall("\\$\\{(.*?)\\}", withouttable[0])
    for paragraph in doc.paragraphs:
        adjust_pattern  = re.findall("\\{{ADJUST\\:(.*?)\\}}", paragraph.text)
        if len(adjust_pattern) > 0 :
            matched_patterns = re.findall("\\$\\{(.*?)\\}", adjust_pattern[0])
            format_type = re.findall("\\((.*?)\\)",adjust_pattern[0])[0]
            format_type = format_type.split(',')
            date_value = attach_field_values(matched_patterns[0],data_dict)
            separate_date = date_value.split('-')
            datefield = separate_date[2][:2]
            value = date(int(separate_date[0]), int(separate_date[1]), int(datefield))
            value = value + datetime.timedelta(int(format_type[1])*365/12)
            value = value + datetime.timedelta(days=int(format_type[0]))
            value = value + datetime.timedelta(int(format_type[2])*365)
            paragraph.text = str(value)
            target_stream = StringIO()
        matched_patterns = re.findall("\\$\\{(.*?)\\}", paragraph.text)
        function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", paragraph.text)
        if len(matched_patterns) > 0 and matched_patterns[0] not in child_obj_metadata:
            for value in matched_patterns :
                            text_in_cell = paragraph.text
                            field_value = attach_field_values(value,data_dict)
                            field_value = text_in_cell.replace('${'+value+'}',str(field_value))
                            paragraph.text = field_value
                            target_stream = StringIO()
        if len(function_list) > 0 :
            field_value = ''
            field_value = generate_functions(function_list,data_dict)
            paragraph.text = field_value
        target_stream = StringIO()
        # doc.save(doc)
    
    para_table_obj = []
    child_tbl_objs = set()
     #Bind values outside child table
    for paragraph in doc.paragraphs:
        child_obj_metadata = re.findall("\\$\\{(.*?)\\}", paragraph.text)
        if len(child_obj_metadata) > 0 :
            child_obj_fields = child_obj_metadata[0].split('.')
            child_tbl_objs.add(child_obj_fields[1])
    
    #Bind SUM values outside the cild table
    for cell in doc.paragraphs:
        count_func_list = re.findall("\\{\\{RowCount:(.*?)\\}}", cell.text)
        has_sum_func = re.findall("\\SUM\\{(.*?)\\}", cell.text)
        if len(count_func_list) > 0 :
            for value in count_func_list :
                    text_in_cell = cell.text
                    field_value = str(len(data_dict[count_func_list[0]]['records']))
                    field_value = text_in_cell.replace('{{RowCount:'+value+'}}',field_value)
                    cell.text = field_value
        elif len(has_sum_func) > 0 :
            print("has_sum_func-->{}".format(has_sum_func))
            format_type = re.findall("(#[A-Z]*)",has_sum_func[0])
            splited_list = has_sum_func[0].split('.')
            if len(format_type) > 0 :
                corrected_field = ''
                if len(format_type) > 0 :
                    corrected_field = splited_list[-1].replace(format_type[0],'').rstrip()
                else :
                    corrected_field = splited_list[-1]
                sum_of_field = 0 
                formatted_type_data = ''
                print('data_dict[splited_list[1]]',splited_list[1])
                for field in data_dict[splited_list[1]]['records'] :
                    print('exists',field)
                    if len(splited_list) > 0 :
                        if len(splited_list) == 3 :
                            try :
                                formatted_type_data = str(field[corrected_field])
                            except:
                                formatted_type_data = ''
                            formatted_type = str(formatted_type_data)
                        elif len(splited_list) == 4:
                            obj_name_match = re.split('Id',splited_list[2])
                            try :
                                formatted_type_data = field[obj_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]].keys() else ''
                            except : 
                                formatted_type_data = ''
                            formatted_type = str(formatted_type_data)
                        elif len(splited_list) == 5 :
                            obj_name_match = re.split('Id',splited_list[2])
                            field_name_match = re.split('Id',splited_list[3])
                            try :
                                formatted_type_data = field[obj_name_match[0]][field_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]][field_name_match].keys() else ''
                            except :
                                formatted_type_data = ''
                            formatted_type = str(formatted_type_data)
                    print(formatted_type_data,type(formatted_type_data))
                    sum_of_field = sum_of_field + float(formatted_type_data)
                    # sum_of_field = sum_of_field + float(field[corrected_field])
                text_in_cell = cell.text
                curr_value= locale.currency(sum_of_field, grouping=True)
                value = curr_value
                field_value = text_in_cell.replace('$SUM{'+has_sum_func[0]+'}',value)
                cell.text = field_value
            else:
                corrected_field = ''
                if len(format_type) > 0 :
                    corrected_field = splited_list[-1].replace(format_type[0],'').rstrip()
                else :
                    corrected_field = splited_list[-1]
                sum_of_field = 0 
                formatted_type_data = ''
                for field in data_dict[splited_list[1]]['records'] :
                    if len(splited_list) > 0 :
                        if len(splited_list) == 3 :
                            print("celltext-->{}".format(corrected_field))
                            print(field)
                            formatted_type = str(field[corrected_field])
                        elif len(splited_list) == 4:
                            obj_name_match = re.split('Id',splited_list[2])
                            formatted_type_data = field[obj_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]].keys() else ''
                            formatted_type = str(formatted_type_data)
                        elif len(splited_list) == 5 :
                            obj_name_match = re.split('Id',splited_list[2])
                            field_name_match = re.split('Id',splited_list[3])
                            try :
                                formatted_type_data = field[obj_name_match[0]][field_name_match][corrected_field] if corrected_field in field[obj_name_match[0]][field_name_match].keys() else ''
                            except :
                                formatted_type_data = ''
                            formatted_type = str(formatted_type_data)
                    sum_of_field = sum_of_field + float(formatted_type)
                text_in_cell = cell.text
                field_value = str(sum_of_field)
                field_value = text_in_cell.replace('$SUM{'+has_sum_func[0]+'}',field_value)
                cell.text = field_value
    table_obj_to_bind_list = []
    for objects in child_tbl_objs :
         fields_list = []
         for paragraph in doc.paragraphs:
             child_obj_metadata = re.findall("\\$\\{(.*?)\\}", paragraph.text)
             if len(child_obj_metadata) > 0 :
                 child_obj_fields = child_obj_metadata[0].split('.')
                 if objects == child_obj_fields[1] :
                     fields_list.append(child_obj_metadata[0])
         table_obj_to_bind_list.append({'objName':objects,'fieldList':fields_list})
    print("table_obj_to_bind_list-->",table_obj_to_bind_list)

    for just_iterate in table_obj_to_bind_list :
        for paragraph in doc.paragraphs:
            table_obj = re.findall("\\$tbl{START:(.*):", paragraph.text)
            if len(table_obj) == 0:
                table_obj = re.findall("\\$tbl\\{START:(.*?)\\}", paragraph.text)
            if len(table_obj) > 0 :
                if just_iterate['objName'] == table_obj[0]:
                    for record in data_dict[table_obj[0]]['records'] :
                        for fields in just_iterate['fieldList'] :
                            child_obj_fields = fields.split('.')
                            field_pattern = fields.split('.')
                            format_type = re.findall("(#[A-Z]*)",fields)
                            corrected_field = ''
                            if len(format_type) > 0 :
                                corrected_field = field_pattern[-1].replace(format_type[0],'').rstrip()
                            else :
                                corrected_field = field_pattern[-1]
                                splited_list = corrected_field.split('.')
                            if len(field_pattern) == 3 :
                                try :
                                    formatted_type_data = record[corrected_field]
                                except :
                                    formatted_type_data = ''
                                formatted_type = str(formatted_type_data)
                            elif len(field_pattern) == 4:
                                obj_name_match = re.split('Id',field_pattern[2])
                                try :
                                    formatted_type_data = record[obj_name_match[0]][corrected_field] if corrected_field in record[obj_name_match[0]].keys() else ''
                                except :
                                    formatted_type_data = ''
                                formatted_type = str(formatted_type_data)
                            elif len(field_pattern) == 5 :
                                obj_name_match = re.split('Id',field_pattern[2])
                                field_name_match = re.split('Id',field_pattern[3])
                                try :
                                    formatted_type_data = record[obj_name_match[0]][field_name_match[0]][corrected_field] if corrected_field in record[obj_name_match[0]][field_name_match[0]].keys() else ''
                                except :
                                    formatted_type_data = ''
                                formatted_type = str(formatted_type_data)
                            if len(format_type) > 0 and format_type[0] == '#NUMBER' :
                                value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                            elif len(format_type) > 0 and format_type[0] == '#CURRENCY' :
                                    curr_value= locale.currency(formatted_type_data, grouping=True)
                                    value = curr_value
                            elif len(format_type) > 0 and format_type[0] == '#DATE' :
                                separate_date = formatted_type.split('-')
                                datefield = separate_date[2][:2]
                                value = date(int(separate_date[0]), int(separate_date[1]), int(datefield)).ctime()
                                value = value.split(' ')
                                value = value[1]+' '+value[2]+','+''+value[-1]
                            field_name = value if len(format_type) > 0 else formatted_type
                            print('field_name-->',field_name)
                            paragraph.insert_paragraph_before(field_name)
                            target_stream = StringIO()

    for just_iterate in table_obj_to_bind_list :
        for paragraph in doc.paragraphs:
            table_obj = re.findall("\\$tbl{START:(.*):", paragraph.text)
            table_end = re.findall("\\$tbl\\{END:(.*?)\\}", paragraph.text)
            if len(table_obj) == 0 :
                table_obj = re.findall("\\$tbl\\{START:(.*?)\\}", paragraph.text)
            if len(table_obj) > 0 :
                if just_iterate['objName'] == table_obj[0]:
                   paragraph.text = ''
            elif len(table_end) > 0 :
                if just_iterate['objName'] == table_end[0]:
                    paragraph.text = ''
            else :
                child_obj_metadata = re.findall("\\$\\{(.*?)\\}", paragraph.text)
                if len(child_obj_metadata) > 0 :
                    child_obj_fields = child_obj_metadata[0].split('.')
                    if just_iterate['objName'] == child_obj_fields[1] :
                        paragraph.text = ''
    target_stream = StringIO()


    alltext_in_tbl = []
    for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    alltext_in_tbl.append(cell.text)
    alltext_in_tbl = '\n'.join(alltext_in_tbl)
    table_patterns = re.findall("\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:.*?\\}", alltext_in_tbl.replace('\n', ' ').replace('\r', ''))
    table_pattern_list = get_all_table_patterns(alltext_in_tbl.replace('\n', ' ').replace('\r', ''))
    table_pattern_string = ' '.join(table_pattern_list)
    print('table_pattern_string-->',table_pattern_string)
    for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matched_patterns = re.findall("\\$\\{(.*?)\\}", cell.text)
                    function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", cell.text)
                    field_value = ''
                    if len(function_list) > 0 :
                            field_value = generate_functions(function_list,data_dict)
                            cell.text = field_value
                    elif len(matched_patterns) > 0 :
                        for value in matched_patterns :
                            if len(table_patterns)>0 and matched_patterns[0] in table_pattern_string:
                                print('matched_patterns-->',matched_patterns)
                                pass
                            else:
                                text_in_cell = cell.text
                                field_value = attach_field_values(value,data_dict)
                                field_value = text_in_cell.replace('${'+value+'}',str(field_value))
                                # cell.text = field_value
    
    target_stream = StringIO()
    
    # Iterating tables to bind parent field values
    if len(doc.tables) > 0:
        table_fields_list = []
        alltext_in_doc =[]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    alltext_in_doc.append(cell.text)
                    # cell.text = attach_field_values(cell.text,data_dict,file_path)
                    matched_patterns = re.findall("\\$\\{(.*?)\\}", cell.text)
                    function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", cell.text)
                    
                    if len(matched_patterns) > 0 :
                        if len(table_patterns)>0 and matched_patterns[0] in table_pattern_string:
                                pass
                        else:
                            for value in matched_patterns :
                                    cell.text = attach_field_values(value,data_dict)
                    elif len(function_list) > 0 :
                        field_value = ''
                        field_value = generate_functions(function_list,data_dict)
                        cell.text = field_value
                    target_stream = StringIO()
                    
                    # doc.save(doc)
        alltext_in_doc = '\n'.join(alltext_in_doc)
        table_values = re.findall("\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:.*?\\}", alltext_in_doc.replace('\n', ' ').replace('\r', ''))
        
        if len(table_values) > 0 :
            table_fields_list = re.findall("\\$\\{(.*?)\\}",table_pattern_string)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        table_row_details = re.findall("\\$\\{(.*?)\\}", cell.text)
                        count_func_list = re.findall("\\{\\{RowCount:(.*?)\\}}", cell.text)
                        has_sum_func = re.findall("\\SUM\\{(.*?)\\}", cell.text)
                        if len(count_func_list) > 0 :
                            for value in count_func_list :
                                    text_in_cell = cell.text
                                    field_value = str(len(data_dict[count_func_list[0]]['records']))
                                    field_value = text_in_cell.replace('{{RowCount:'+value+'}}',field_value)
                                    cell.text = field_value
                        elif len(has_sum_func) > 0 :
                            print("has_sum_func-->{}".format(has_sum_func))
                            format_type = re.findall("(#[A-Z]*)",has_sum_func[0])
                            splited_list = has_sum_func[0].split('.')
                            if len(format_type) > 0 :
                                corrected_field = ''
                                if len(format_type) > 0 :
                                    corrected_field = splited_list[-1].replace(format_type[0],'').rstrip()
                                else :
                                    corrected_field = splited_list[-1]
                                sum_of_field = 0 
                                formatted_type_data = ''
                                for field in data_dict[splited_list[1]]['records'] :
                                    if len(splited_list) > 0 :
                                        if len(splited_list) == 3 :
                                            formatted_type = str(field[corrected_field])
                                            formatted_type_data = field[corrected_field]
                                        elif len(splited_list) == 4:
                                            obj_name_match = re.split('Id',splited_list[2])
                                            formatted_type_data = field[obj_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]].keys() else ''
                                            formatted_type = str(formatted_type_data)
                                        elif len(splited_list) == 5 :
                                            obj_name_match = re.split('Id',splited_list[2])
                                            field_name_match = re.split('Id',splited_list[3])
                                            try :
                                                formatted_type_data = field[obj_name_match[0]][field_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]][field_name_match].keys() else ''
                                            except :
                                                formatted_type_data = ''
                                            formatted_type = str(formatted_type_data)
                                    print(formatted_type_data,type(formatted_type_data))
                                    sum_of_field = sum_of_field + float(formatted_type_data)
                                    # sum_of_field = sum_of_field + float(field[corrected_field])
                                text_in_cell = cell.text
                                curr_value= locale.currency(sum_of_field, grouping=True)
                                value = curr_value
                                field_value = text_in_cell.replace('$SUM{'+has_sum_func[0]+'}',value)
                                cell.text = field_value
                            else:
                                corrected_field = ''
                                if len(format_type) > 0 :
                                    corrected_field = splited_list[-1].replace(format_type[0],'').rstrip()
                                else :
                                    corrected_field = splited_list[-1]
                                sum_of_field = 0 
                                formatted_type_data = ''
                                for field in data_dict[splited_list[1]]['records'] :
                                    if len(splited_list) > 0 :
                                        if len(splited_list) == 3 :
                                            try :
                                                formatted_type_data = str(field[corrected_field])
                                            except :
                                                formatted_type_data = ''
                                            formatted_type = str(formatted_type_data)
                                        elif len(splited_list) == 4:
                                            obj_name_match = re.split('Id',splited_list[2])
                                            try :
                                                formatted_type_data = field[obj_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]].keys() else ''
                                            except :
                                                formatted_type_data = ''
                                            formatted_type = str(formatted_type_data)
                                        elif len(splited_list) == 5 :
                                            obj_name_match = re.split('Id',splited_list[2])
                                            field_name_match = re.split('Id',splited_list[3])
                                            try :
                                                formatted_type_data = field[obj_name_match[0]][field_name_match][corrected_field] if corrected_field in field[obj_name_match[0]][field_name_match].keys() else ''
                                            except :
                                                formatted_type_data = ''
                                            formatted_type = str(formatted_type_data)
                                    sum_of_field = sum_of_field + float(formatted_type)
                                text_in_cell = cell.text
                                field_value = str(sum_of_field)
                                field_value = text_in_cell.replace('$SUM{'+has_sum_func[0]+'}',field_value)
                                cell.text = field_value
                        
                        
                        if len(table_row_details) > 0 and table_row_details[0] not in table_fields_list :
                            matched_patterns = re.findall("\\$\\{(.*?)\\}", cell.text)
                            function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", cell.text)
                            field_value = ''
                            if len(function_list) > 0 :
                                    field_value = generate_functions(function_list,data_dict)
                                    cell.text = field_value
                            elif len(matched_patterns) > 0 :
                                table_obj = re.findall("\\$tbl\\{START:(.*?)\\}", cell.text)
                                for value in matched_patterns :
                                        text_in_cell = cell.text
                                        field_value = attach_field_values(value,data_dict)
                                        field_value = text_in_cell.replace('${'+value+'}',field_value)
                                        cell.text = field_value
            target_stream = StringIO()
            # r = requests.post("https://yourInstance.salesforce.com/services/data/v23.0/sobjects/ContentVersion",data=obj_wrapper,headers = 
            # curl https://yourInstance.salesforce.com/services/data/v23.0/sobjects/ContentVersion -H "Authorization: Bearer token" -H "Content-Type: multipart/form-data; boundary=\"boundary_string\"" --data-binary @NewContentVersion.json
            # doc.save(doc)
 
        def remove_row(table, row):
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)          
        
        # Iterating tables to bind child field values       
        for table in doc.tables:
            column_value_list = []
            head_obj = []
            row_to_add = []
            for row_index,row in enumerate(table.rows) : 
                    for column_index,cell in enumerate(row.cells):
                            check_child =  re.findall("\\$\\{(.*?)\\}", cell.text)
                            if len(check_child) > 0 and  check_child[0] in  table_fields_list :
                                table_row_details = re.findall("\\$\\{(.*?)\\}", cell.text)
                            table_obj = re.findall("\\$tbl{START:(.*):", cell.text)
                            if len(table_obj) == 0:
                                table_obj = re.findall("\\$tbl\\{START:(.*?)\\}", cell.text)
                            table_end = re.findall("\\$tbl\\{END:(.*?)\\}", cell.text)
                            if len(table_obj) > 0:
                                row_to_add = table.row_cells(row_index)
                                head_obj = re.findall("\\$tbl{START:(.*):", cell.text)
                                if len(head_obj) == 0:
                                    head_obj = re.findall("\\$tbl\\{START:(.*?)\\}", cell.text)
                                head_obj[0] = head_obj[0].strip()
                                row_columns =[]
                                for cell in row_to_add:
                                    if cell.text not in row_columns :
                                        row_columns.append(cell.text)
                                if len(head_obj) > 0 :  
                                    if len(check_child) > 0 and  check_child[0] in  table_fields_list : 
                                        for i,record in enumerate(data_dict[head_obj[0]]['records']) : 
                                            current_row = table.rows[row_index]
                                            border_copied = copy.deepcopy(current_row._tr)
                                            tr = border_copied
                                            current_row._tr.addnext(tr)
                                            for j,column in enumerate(row_columns):
                                                table_pattern  = re.findall("\\$\\{(.*?)\\}", column)
                                                if len(table_pattern) > 0 :
                                                    field_pattern = table_pattern[0].split('.')
                                                    format_type = re.findall("(#[A-Z]*)",table_pattern[0])
                                                    corrected_field = ''
                                                    if len(format_type) > 0 :
                                                        corrected_field = field_pattern[-1].replace(format_type[0],'').rstrip()
                                                    else :
                                                        corrected_field = field_pattern[-1]
                                                        splited_list = corrected_field.split('.')
                                                    if len(field_pattern) == 3 :
                                                        try : 
                                                            formatted_type_data = record[corrected_field]
                                                        except : 
                                                            formatted_type_data = ''
                                                        formatted_type = str(formatted_type_data)
                                                    elif len(field_pattern) == 4:
                                                        obj_name_match = re.split('Id',field_pattern[2])
                                                        try :
                                                            formatted_type_data = record[obj_name_match[0]][corrected_field] if corrected_field in record[obj_name_match[0]].keys() else ''
                                                        except :
                                                            formatted_type_data = ''
                                                        formatted_type = str(formatted_type_data)
                                                    elif len(field_pattern) == 5 :
                                                        obj_name_match = re.split('Id',field_pattern[2])
                                                        field_name_match = re.split('Id',field_pattern[3])
                                                        try :
                                                            formatted_type_data = record[obj_name_match[0]][field_name_match[0]][corrected_field] if corrected_field in record[obj_name_match[0]][field_name_match[0]].keys() else ''
                                                        except :
                                                            formatted_type_data = ''
                                                        formatted_type = str(formatted_type_data)
                                                    if len(format_type) > 0 and format_type[0] == '#NUMBER' :
                                                        value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                                                    elif len(format_type) > 0 and format_type[0] == '#CURRENCY' :
                                                            curr_value= locale.currency(formatted_type_data, grouping=True)
                                                            value = curr_value
                                                    elif len(format_type) > 0 and format_type[0] == '#DATE' :
                                                        separate_date = formatted_type.split('-')
                                                        datefield = separate_date[2][:2]
                                                        value = date(int(separate_date[0]), int(separate_date[1]), int(datefield)).ctime()
                                                        value = value.split(' ')
                                                        value = value[1]+' '+value[2]+','+''+value[-1]
                                                    field_name = value if len(format_type) > 0 else formatted_type
                                                    table.cell(row_index+1, j).text = field_name
                                                    # table.rows[row_index+1].height_rule = WD_ROW_HEIGHT.AUTO
                                                    # print("table-->{}".format(table.rows[row_index+1].height_rule))
                                                    table.rows[row_index+1].height = 1        
                            if len(table_end) > 0 :
                                remove_row(table, table.rows[row_index])
        target_stream = StringIO()
        # doc.save(doc) 
    
#Method to get field index
#Parameters (fieldName, metaData, objName)
def get_field_index(field_name, data,list_name,obj_name):
        obj_list = list()
        for record in data[list_name]:
            obj_list.append(record[obj_name])
        if len(obj_list) > 0:
            try:
                return obj_list.index(field_name)
            except ValueError:
                return -1
        else: 
            return -1     

#Method to manipulate functions in the document
def generate_functions(function_list,data_dict) :
    if_condition_list = re.findall("IF\\((.*?)\\)", function_list[0])
    if len(if_condition_list) > 0 :
        conditon_value,true_value,false_value = if_condition_list[0].split(',')[0],if_condition_list[0].split(',')[1],if_condition_list[0].split(',')[2]
        field_name_list = re.findall("\\$\\{(.*?)\\}", conditon_value)
        if '==' in str(conditon_value):
            conv_value_to_str = re.split('== ',str(conditon_value))
        if '!=' in str(conditon_value):
            conv_value_to_str = re.split('!= ',str(conditon_value))
        if '>=' in str(conditon_value):
            conv_value_to_str = re.split('>= ',str(conditon_value))
        if '<=' in str(conditon_value):
            conv_value_to_str = re.split('<= ',str(conditon_value))
        if '>' in str(conditon_value):
             conv_value_to_str = re.split('> ',str(conditon_value))
        if '<' in str(conditon_value):
            conv_value_to_str = re.split('< ',str(conditon_value))
        # print("conv_value_to_str-->",conditon_value)
        print("addstring-->","'"+conv_value_to_str[-1]+"'")
        added_changes = conditon_value.replace(conv_value_to_str[-1],"'"+conv_value_to_str[-1]+"'")
        print("added_changes-->",added_changes)
        if len(field_name_list) > 0 :
                splited_list = field_name_list[0].split('.')
                if len(splited_list) == 2 :
                    field_value = data_dict[splited_list[1]]
                elif len(splited_list) == 3 :
                    obj_name_match = re.split('Id',splited_list[1])
                    field_value = data_dict[obj_name_match[0]][splited_list[2]]
                elif len(splited_list) == 4 :
                    parent_name_match = re.split('Id',splited_list[1])
                    grand_name_match = re.split('Id',splited_list[2])
                    field_value = data_dict[parent_name_match[0]][grand_name_match[0]][splited_list[3]]
        field_value = str(field_value)
        # field_value = field_value.replace(" ","")
        val = added_changes.replace('${'+field_name_list[0]+'}',"'"+field_value.strip()+"'")
        
        print("evalBefore-->","true_value if "+val+" else false_value")
        print("regex",bool(re.match('^(?=.*[a-zA-Z])',val)))
        if bool(re.match('^(?=.*[a-zA-Z])',val)) == False:
            val = val.replace("'","")    
        cons = eval("true_value if "+val+" else false_value")
        print("Result-->",cons)
        return cons
    else : 
        return "Error"


#To bind values from salesforce to the matched string
#Parameters(fieldName, metaData, filePath)
def attach_field_values(obj_to_bind,data_dict) :
     function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", obj_to_bind)
     field_name = ''
     if len(function_list) > 0 :
            field_name = generate_functions(function_list,data_dict)
     else :
            format_type = re.findall("(#[A-Z]*)",obj_to_bind)
            corrected_field = ''
            if len(format_type) > 0 :
                corrected_field = obj_to_bind.replace(format_type[0],'').rstrip()
            else :
                corrected_field = obj_to_bind
            splited_list = corrected_field.split('.')

            if len(splited_list) == 2 :
                formatted_type_data = data_dict[splited_list[1]] if splited_list[1] in data_dict.keys() else ''
                formatted_type = str(formatted_type_data)
                if len(format_type) > 0 and format_type[0] == '#NUMBER' :
                    value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                elif len(format_type) > 0 and format_type[0] == '#CURRENCY' :
                    curr_value= locale.currency(formatted_type_data, grouping=True)
                    value = curr_value
                    
                elif len(format_type) > 0 and format_type[0] == '#DATE' :
                    separate_date = formatted_type.split('-')
                    datefield = separate_date[2][:2]
                    value = date(int(separate_date[0]), int(separate_date[1]), int(datefield)).ctime()
                    value = value.split(' ')
                    value = value[1]+' '+value[2]+','+''+value[-1]
                field_name = value if len(format_type) > 0 else formatted_type
            elif len(splited_list) == 3 :
                obj_name_match = re.split('Id',splited_list[1])
                try :
                    print('object-->',obj_name_match[0])
                    print('field-->',[splited_list[2]])
                    print('data_dict[obj_name_match[0]].keys()-->',data_dict[obj_name_match[0]].keys())
                    formatted_type_data = data_dict[obj_name_match[0]][splited_list[2]] if splited_list[2] in data_dict[obj_name_match[0]].keys() else ''
                except KeyError:
                    formatted_type_data = ''
                print(formatted_type_data)
                formatted_type = str(formatted_type_data)
                if len(format_type) > 0 and format_type[0] == '#NUMBER' :
                    value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                elif len(format_type) > 0 and format_type[0] == '#CURRENCY' :
                    print('formatted_type_data-->',formatted_type_data)
                    curr_value= locale.currency(formatted_type_data, grouping=True)
                    value = curr_value
                elif len(format_type) > 0 and format_type[0] == '#DATE' :
                    separate_date = formatted_type.split('-')
                    datefield = separate_date[2][:2]
                    value = date(int(separate_date[0]), int(separate_date[1]), int(datefield)).ctime()
                    value = value.split(' ')
                    value = value[1]+' '+value[2]+','+''+value[-1]
                field_name = value if len(format_type) > 0 else formatted_type
            elif len(splited_list) == 4 :
                 obj_name_match = re.split('Id',splited_list[1])
                 obj_field_name = re.split('Id',splited_list[2])
                 try :
                    formatted_type_data = data_dict[obj_name_match[0]][obj_field_name[0]][splited_list[3]]
                 except KeyError:
                     formatted_type_data = ''
                 formatted_type = str(formatted_type_data)
                #  formatted_type = data_dict[parent_name_match[0]][grand_name_match[0]][splited_list[3]] if [splited_list[3]] in parent_list.keys() else ''
                 if len(format_type) > 0 and format_type[0] == '#NUMBER' :
                    value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                 elif len(format_type) > 0 and format_type[0] == '#CURRENCY' :
                    curr_value= locale.currency(formatted_type_data, grouping=True)
                    value = curr_value
                 elif len(format_type) > 0 and format_type[0] == '#DATE' :
                    separate_date = formatted_type.split('-')
                    datefield = separate_date[2][:2]
                    value = date(int(separate_date[0]), int(separate_date[1]), int(datefield)).ctime()
                    value = value.split(' ')
                    value = value[1]+' '+value[2]+','+''+value[-1]
                 field_name = value if len(format_type) > 0 else formatted_type

     return field_name

if __name__ == "__main__":
    app.run(ssl_context="adhoc")
    



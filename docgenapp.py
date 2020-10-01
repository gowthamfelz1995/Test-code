import os
from flask import Flask, request, render_template, jsonify, url_for, redirect
from oauthlib.oauth2 import WebApplicationClient
import requests 
from werkzeug.utils import secure_filename 
from xml.etree import ElementTree
from docx import Document
from docx.document import Document as _Document
import docx
from base64 import b64decode
import re
import json  
from datetime import date
import datetime
import copy
import xml.etree.ElementTree as ET
from docx.enum.table import WD_ROW_HEIGHT
from Models.objwrapper import obj_wrap
from Models.fieldwrap import field_wrap_obj
from Models.grandwrap import grand_wrap_obj
from Models.parentwrap import parent_wrap_obj
from Models.childwrap import child_wrap_obj
from Models.user import user_credentials
from _io import BytesIO, StringIO
import base64
import io
import random
import string
import locale
locale.setlocale(locale.LC_ALL, 'en_US.utf-8')
from sqlalchemy import create_engine,Column,Integer,String,Date,ForeignKey,MetaData,Table
from sqlalchemy.ext.declarative import declarative_base 
import pymysql.cursors
from _datetime import date, timedelta
import uuid
import shutil
from flask_cors import CORS
import jwt
from oauthlib.common import urlencode
import pymongo
from bson.objectid import ObjectId
from boxsdk import JWTAuth

auth = JWTAuth(
    client_id='fj5klf35gaydgnx2hewpfdg4uqz85ri7',
    client_secret='mM2d0ZScYgg6dE22jShBgK6y5xL6xFID',
    enterprise_id='635964596',
    jwt_key_id='061jkz4x',
    rsa_private_key_file_sys_path=None,
    rsa_private_key_data = '-----BEGIN ENCRYPTED PRIVATE KEY-----\nMIIFDjBABgkqhkiG9w0BBQ0wMzAbBgkqhkiG9w0BBQwwDgQICUFlpOqJ2DkCAggA\nMBQGCCqGSIb3DQMHBAhk2ZclZTqcNASCBMhSLnADIn+MhR1PtS0YWLej0bUrIxjE\n7uzrtYwXN9b8G7Y7PUhroEetlRIS5Uf4qEDsYFJeZwlW/wd9Xzzjdhg+bUMmVLC5\n57tcPnS9Sv0Q3gDcGiJrhjm9k7u17/GgHhCOeUeakZX8U/RJfasijw5X8ue8+esY\ng+99BB83H+9anSjTYqic2dU7PtwDdx9AYeru6g1psBjZurzT0eE1SWdTOw5QGhG4\n78+gBsuzuU4D2SefS1wMFiuY45JvxUq6+zAGtzRi8MYXYruRsi3vQRkctTqcJZil\n6g7yxs/oY0xaSNocMY7y8kuHVirdYzj3KZX/+p8EvVB+1pzVJ50dFtJntzdpcu03\nfCfLYz2Vod6TXnGs2Dr6Zzq9NSb5otumvn6FrD5QeAoVKkEcakPbGOLACBKCJe7B\nfoDB2FJXC2LqZwDsf0X7kLcJh7xaf6l2HOOp73wsZUfpifGmmswmL3dfEbZ3uLeR\nIBIdSlNyWuwrOuNc+5fl0V1zBdys3RGo4pavwq7QDFVSg3FgwtAqfOFICUlO48EL\nuF9Z3xp8pz+rlfYiF4iMpbIqt1/OD6GOFH2PiVIUaf4xB7Arbbbv93JKnTzwj8j1\nbKowjITWVc4RmPI0XbCVvukAYm4hY1ohlTaFVs5+3Dg6BQTqEtkfxknCphtPUoGD\nKV+c0YxI4jz9bYLPAWsMvwlvfYg10kUleyzKZAGt+NvBeplbi73Je7as1ZcvNcJX\nJ1jBLatH24RB7XU0N2r0KjIKMaE7pMDghTjtoJ1DczIi2t+J5dymwqAyoba0FktF\nA+T/AsAfjs95ashgekL0g0odKeCQL6YOnb41AKFZ9m6aOOONF0q8S4oLKYdqoL5U\nWq/2DnMyiKNH7eZfbAKMAJzYDjjkVKXj2MBDg/tF4ibwoVoa6jJZT6eKi82moCr5\n3LPQpJsIlfQ5nZg6faito35KlKEUCLQkSPP9p+sAZ3qyqirvkzps1c2FRyuUnVWH\nCMPUKGHcOhdq/2ziSLS8a1X95PxpWHCOS4BRnHEiGX0m2WleqkiQsJRnXrAw8Ft8\neu6knlL1o51EE4EaQeipfHaBxRDhdLMxEJGmj4mOTOP9Q8It9Sk4FNu8Sf3bXV85\nBpK30tBHFgSo04DQdxFrcHykEuFR0CjEf26us6Thp8YkAov/ogQb6EdjaYUQpm2P\nGdWxle14L/QAhkzXGA3cVng5bbeqU3gdagCe+UN1gCWGXer6OFmA3jjy5jZCrlDM\nd1O2HonoFdlh0V+RhJtvJWi2P/KcIOu/LroBnR0wcOles/TqnSHZ3LvPp2f+SfP+\nL2KL+jgZR24vofh+owEELYSD2luoi1H7X79cCWP6qxfYBj7jKzUBUkVjKOZ5QwIo\nlpCspT9mHzDB+K83p9nwwrnYy+J55InT0tiLtWvpN3og6kPy+vdX9SN1rp0VImxO\n6n1pfzYwfbhGLgc953/avlU9vSSHjAA9wFsONB344H+QMT1pgMh6D4O42XfZzriY\ne/rgz/EH7v39RQL4Y4VJbmdv9WW3NqBfLGwGgVuGFhh0wSc3ct/PYImeiijYviiD\nOvzIFZexsc5jgms0qr7Iby2fqYAwWYhxGN/uqpfDq0p9TYti7XUNHsDa0jAQrEb/\no1A=\n-----END ENCRYPTED PRIVATE KEY-----\n',
    rsa_private_key_passphrase='b8560c4b0187d6d5a068a637cfd034a9'
)

access_token = auth.authenticate_instance()

from boxsdk import Client

boxClient = Client(auth)





# connect to MongoDB, change the << MONGODB URL >> to reflect your own connection string
client = pymongo.MongoClient("mongodb://docgen:Aspi2018@cluster0-shard-00-00.6swxi.mongodb.net:27017,cluster0-shard-00-01.6swxi.mongodb.net:27017,cluster0-shard-00-02.6swxi.mongodb.net:27017/sample_analytics?ssl=true&replicaSet=atlas-lo4ypi-shard-0&authSource=admin&retryWrites=true&w=majority")
db = client.sample_analytics




UPLOAD_FOLDER = str(os.getcwd())+'/python-mailmerge/Document'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif'}
path = str(os.getcwd())+'/Document'
app = Flask(__name__)
cors = CORS(app, resources={"/*": {"origins": "*"}})
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER 
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
file_path = ''
document_data = ''
document_data_initial = ''
record_id = ''
file_name = ''
file_type = ''
table_pattern_list = []
folder_id_dyn = ''



@app.route('/generate_file', methods =['POST'])
def generate_document():
    # content = request.files['file'].read()
    file_id = request.headers['file_id']
    record_id = request.headers['record_id']
    file_content = boxClient.file(file_id=file_id).content()
    # file_name = request.files['file'].filename
    file_type = 'DOCX'
    file_name = 'test'+'.docx'
    bytes = b64decode(file_content)
    source_stream = BytesIO(file_content)
    doc = Document(source_stream)
    source_stream.close()
    full_text = []
    field_list = []
    child_obj_metadata = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    document_data = '\n'.join(full_text)
    field_list = re.findall("\\$\\{(.*?)\\}", document_data)
    child_obj_metadata = re.findall("\\$tbl\\{(.*?)\\}", document_data)
    withouttable = get_all_table_patterns(document_data.replace('\n', ' ').replace('\r', ''))
    if len(withouttable) > 0 : 
        child_obj_metadata = re.findall("\\$\\{(.*?)\\}", withouttable[0])
        field_list = list(set(field_list) - set(child_obj_metadata))
    head_child_obj = '' 
    alltext_in_tbl = []
    for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    alltext_in_tbl.append(cell.text)
    alltext_in_tbl = '\n'.join(alltext_in_tbl)
    check_value_str = re.findall("\\$tbl\\{(.*?)\\}", alltext_in_tbl)
    table_patterns = re.findall("\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:.*?\\}", alltext_in_tbl.replace('\n', ' ').replace('\r', ''))
    
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                fields_in_cell = re.findall("\\$\\{(.*?)\\}",cell.text)
                child_obj_values = re.findall("\\$tbl\\{(.*?)\\}",cell.text)
                child_obj_str = []
                if len(child_obj_values) > 0:
                   child_obj_str =  re.findall("\\$tbl\\{START:(.*?)\\}", cell.text) 
                if len(child_obj_str) > 0:
                   head_child_obj = child_obj_str[0]
                    
                # if len(fields_in_cell) > 0 and fields_in_cell[0].split('.')[1] == head_child_obj.strip():
                #     print("EXISTS")
                #     for field in fields_in_cell :
                #         child_obj_metadata.append(field)
                # elif len(fields_in_cell) > 0 :
                #     for field in fields_in_cell :
                #         field_list.append(field)
                if len(fields_in_cell) > 0 :
                    for field in fields_in_cell :
                        parent_obj_api = head_child_obj.split('(')
                        obj_name = parent_obj_api[1].split(')')[0]
                        field_name = parent_obj_api[0]
                        if field.split('.')[1] == field_name:
                            child_obj_metadata.append(field)
                        else:
                            field_list.append(field)
    field_list = list(dict.fromkeys(field_list))
    
    
    #This block will do json formation of fields and parent objects which will be returned to salesforce for retrieving data             
    if len(field_list) > 0 :
        obj_wrapper = obj_wrap(field_list[0].split('.')[0],False,[],[],[])
        field_wrapper = []
        parent_wrapper = []
        parent_field_wrapper = []
        grand_parent_field_wrapper = []
        grand_wrapper = []
        for field in field_list:
            format_type = re.findall("(#[A-Z]*)",field)
            if len(format_type) > 0 :
                corrected_field = field.replace(format_type[0].strip(),'')
            else :
                corrected_field = field
            parent_obj = corrected_field.split('.')
            if len(parent_obj) == 2:
                field_wrap = field_wrap_obj(parent_obj[-1],False)
                if field_wrap not in field_wrapper:
                    field_wrapper.append(field_wrap)

            else:
                field_wrap = field_wrap_obj(parent_obj[1],False)
                if field_wrap.__dict__ not in field_wrapper:
                    field_wrapper.append(field_wrap.__dict__)
                filtered_obj = parent_obj[1:len(parent_obj)]

                if len(filtered_obj) == 3:
                    parent_field_wrap = field_wrap_obj(filtered_obj[1],False)
                    if parent_field_wrap not in parent_field_wrapper:
                        parent_field_wrapper.append(parent_field_wrap)
                    grand_parent_field_wrap = field_wrap_obj(filtered_obj[2],False)
                    if grand_parent_field_wrap not in grand_parent_field_wrapper:
                        grand_parent_field_wrapper.append(grand_parent_field_wrap)
                    grand_wrap = grand_wrap_obj(filtered_obj[1],False,grand_parent_field_wrapper)
                    if len(grand_wrapper) > 0:
                        check_grobj_list = list()
                        for obj in grand_wrapper:
                            check_grobj_list.append(obj.objName)

                        if grand_wrap.objName not in check_grobj_list:
                            grand_wrapper.append(grand_wrap)

                        elif {
                                'fieldName': filtered_obj[2]
                        } not in grand_wrapper[check_grobj_list.index(
                                grand_wrap.objName)].fieldWrapperList:
                            grand_wrapper[check_grobj_list.index(
                                grand_wrap.objName)].fieldWrapperList.append(
                                    field_wrap_obj(filtered_obj[2],False))
                    else:
                        grand_wrapper.append(grand_wrap)
                    parent_wrap = parent_wrap_obj(filtered_obj[0],False,parent_field_wrapper,[],[],grand_wrapper)

                    if len(parent_wrapper) > 0:
                        check_obj_list = list()
                        for obj in parent_wrapper:
                            check_obj_list.append(obj.objName)
                        if parent_wrap.objName not in check_obj_list:
                            parent_wrapper.append(parent_wrap)

                        else:

                            if field_wrap_obj(filtered_obj[1],False) not in parent_wrapper[check_obj_list.index(
                                    parent_wrap.objName)].fieldWrapperList:
                                parent_wrapper[check_obj_list.index(
                                    parent_wrap.objName
                                )].fieldWrapperList.append(field_wrap_obj(filtered_obj[1],False))
                                # if 'grandObjWrapperList' in parent_wrapper[
                                #         check_obj_list.index(
                                #             filtered_obj[0])]:
                                parent_wrapper[check_obj_list.index(
                                    parent_wrap.objName
                                )].grandObjWrapperList.append(grand_wrap_obj(filtered_obj[1],False,[field_wrap_obj(filtered_obj[2],False)]))

                                # else:
                                # parent_wrapper[check_obj_list.index(
                                #     parent_wrap.objName
                                # )].grandObjWrapperList = [grand_wrap_obj(filtered_obj[1],False,[field_wrap_obj(filtered_obj[2],False)])]
                                
                                
                                
                                

                            else:

                                check_grandobj_list = list()
                                if 'grandObjWrapperList' in parent_wrapper[
                                        check_obj_list.index(
                                            parent_wrap.objName)]:

                                    for obj in parent_wrapper[check_obj_list.index(
                                            parent_wrap.objName
                                    )].grandObjWrapperList:
                                        check_grandobj_list.append(obj.objName)
                                else:
                                    check_grandobj_list = []
                                if grand_wrap.objName not in check_grandobj_list:
                                    grand_wrapper.append(grand_wrap)

                                else:

                                    if field_wrap_obj(filtered_obj[2],False) not in parent_wrapper[check_obj_list.index(
                                            parent_wrap.objName
                                    )].grandObjWrapperList[
                                            check_grandobj_list.index(
                                                grand_wrap.objName
                                            )].fieldWrapperList:

                                        parent_wrapper[check_obj_list.index(
                                            parent_wrap.objName
                                        )].grandObjWrapperList[
                                            check_grandobj_list.index(
                                                grand_wrap.objName
                                            )].fieldWrapperList.append(field_wrap_obj(filtered_obj[2],False))

                    else:
                        parent_wrapper.append(parent_wrap)

                    grand_wrap = {}
                    grand_parent_field_wrap = {}
                    parent_field_wrap = {}
                    parent_wrap = {}
                    parent_field_wrapper = []
                    grand_parent_field_wrapper = []
                    grand_wrapper = []

                elif len(filtered_obj) == 2:
                    parent_field_wrap = field_wrap_obj(filtered_obj[-1],False)
                    check_parent_obj_list = list()
                    for obj in parent_wrapper:
                        check_parent_obj_list.append(obj.objName)
                    if ({
                            'objName': filtered_obj[0]
                    } in check_parent_obj_list
                        ) and parent_field_wrap not in parent_wrapper[
                            check_parent_obj_list.index(
                                filtered_obj[0])]['fieldWrapperList']:
                        parent_field_wrapper.append(parent_field_wrap)

                    parent_wrap = parent_wrap_obj(filtered_obj[0],False,parent_field_wrapper,[],[],[])
                    if len(parent_wrapper) > 0:
                        check_obj_list = list()
                        for obj in parent_wrapper:
                            check_obj_list.append(obj.objName)
                        if parent_wrap.objName not in check_obj_list:
                            parent_wrap.fieldWrapperList = [field_wrap_obj(filtered_obj[-1],False)]
                            parent_wrapper.append(parent_wrap)
                            parent_field_wrapper = []

                        elif {
                                'fieldName': filtered_obj[-1],
                                'isExists': False
                        } not in parent_wrapper[check_obj_list.index(
                                parent_wrap.objName)].fieldWrapperList:
                            parent_wrapper[check_obj_list.index(
                                parent_wrap.objName
                            )].fieldWrapperList.append(field_wrap_obj(filtered_obj[-1],False))
                            parent_field_wrapper = []
                    else:
                        parent_wrap.fieldWrapperList.append(field_wrap_obj(filtered_obj[-1],False))
                        parent_wrapper.append(parent_wrap)
                        parent_field_wrapper = []
                parent_field_wrapper = []
        
        obj_wrapper.fieldWrapperList = field_wrapper
        obj_wrapper.parentObjWrapperList = parent_wrapper
        
        
       
        
        
    parent_wrapper = [] 
    parent_field_wrapper = []
    parent_field_wrap = {}
    parent_wrap = {}
    old_child_obj_meta = []
    
        
    
        
    
    #Method to return index of the obj 
    def check_obj_present(current_obj, child_wrapper):
        obj_list = list()
        for record in child_wrapper.parentObjWrapperList:
            obj_list.append(record.objName)
        if len(obj_list) > 0:
            try:
                return obj_list.index(current_obj)
            except ValueError:
                return -1
        else:
            return -1
    
    #Method to return index of the parent obj
    def check_grand_obj_present(current_obj, child_wrapper, parent_index):
        obj_list = list()

        for record in child_wrapper.parentObjWrapperList[parent_index].grandObjWrapperList:
            obj_list.append(record.objName)
        if len(obj_list) > 0:
            try:
                return obj_list.index(current_obj)
            except ValueError:
                return -1
        else:
            return -1
    #Method to return index of the field obj
    def check_field_obj_present(current_obj, child_wrapper):
        field_list = list()

        for record in child_wrapper.fieldWrapperList:
            field_list.append(record.fieldName)
        if len(field_list) > 0:
            try:
                return field_list.index(current_obj)
            except ValueError:
                return -1
        else:
            return -1
    
    #Method to generate childWrapper
    def generate_child_obj(child_obj, child_wrapper):
        child1_field_wrap = field_wrap_obj(child_obj[2],False)
        field_index = check_field_obj_present(child_obj[2], child_wrapper)
        if len(child_obj) == 3:
            if field_index == -1:
                    child_wrapper.fieldWrapperList.append(child1_field_wrap)
        if len(child_obj) == 4:
            if field_index == -1:
                child_wrapper.fieldWrapperList.append(child1_field_wrap)
            index_value = check_obj_present(child_obj[2], child_wrapper)
            if index_value == -1:
                child_wrapper.parentObjWrapperList.append(parent_wrap_obj(child_obj[2],False,[field_wrap_obj(child_obj[3],False)],[],[],[]))
            else:
                
                if field_wrap_obj(child_obj[3],False) not in child_wrapper.parentObjWrapperList[check_obj_present(child_obj[2],child_wrapper)].fieldWrapperList:
                    child_wrapper.parentObjWrapperList[check_obj_present(child_obj[2],child_wrapper)].fieldWrapperList.append(field_wrap_obj(child_obj[3],False))
        if len(child_obj) == 5:
            if field_index == -1:
                child_wrapper.fieldWrapperList.append(child1_field_wrap)
            index_value = check_obj_present(child_obj[2], child_wrapper)
            if index_value == -1:
                child_wrapper.parentObjWrapperList.append(parent_wrap_obj(child_obj[2],False,
                [field_wrap_obj(child_obj[3],False)],
                [grand_wrap_obj(child_obj[3],False,[field_wrap_obj(child_obj[4],False)])]
                ))
            else:
                if field_wrap_obj(child_obj[3],False) not in child_wrapper.parentObjWrapperList[index_value].fieldWrapperList:
                    child_wrapper.parentObjWrapperList[index_value].fieldWrapperList.append(field_wrap_obj(child_obj[3],False))
                grand_index_value = check_grand_obj_present(child_obj[3], child_wrapper, index_value)
                if grand_index_value == -1:
                    child_wrapper.parentObjWrapperList[index_value].grandObjWrapperList.append(grand_wrap_obj(child_obj[3],False,[field_wrap_obj(child_obj[4],False)]))
                else:
                    child_wrapper.parentObjWrapperList[index_value].grandObjWrapperList[grand_index_value].fieldWrapperList.append(field_wrap_obj(child_obj[4],False))
        return child_wrapper
    
    
            
    if len(child_obj_metadata) > 0:
        child_wrapper = child_wrap_obj('',False,[],[],'')
        child_obj_list = []
        child_wrapper_list = []
        get_condition = []
        check_whr_condition = re.findall("\\$tbl\\{END:(.*?)\\}",alltext_in_tbl)
        cnd_obj_set = dict()
        main_obj = ''
        if len(check_whr_condition) > 0 :
               for cnd_value in check_whr_condition :
                    try:
                        condition_on_table_idx = cnd_value.index('#CND')
                    except ValueError:
                        condition_on_table_idx = -1
                    if condition_on_table_idx != -1 :
                       cnd_table = cnd_value[condition_on_table_idx+4:]
                       cnd_obj = cnd_value[:condition_on_table_idx-1]
                       cnd_obj_set[cnd_obj]=cnd_table
                       
        for field in child_obj_metadata:
            format_type = re.findall("(#[A-Z]*)",field)
            if len(format_type) > 0 :
                corrected_field = field.replace(format_type[0],'')
            else :
                corrected_field = field
            child_obj = corrected_field.split('.')
            main_obj = child_obj[0]
            if child_obj not in old_child_obj_meta:
                if child_obj[1] not in child_obj_list:
                    head_obj = re.findall("\\$tbl{START:[A-Za-z]\\:(.*)",document_data)
                    braces_text = re.findall("\\$tbl\\{(.*?)\\}", alltext_in_tbl)
                    selected_string = ''
                    for text_string in braces_text:
                        if child_obj[1] in text_string:
                            selected_string = text_string
                            break
                    object_name_with_field = selected_string.split(':')[1]
                    if len(check_whr_condition) > 0 :
                        try :
                            child_obj_check = child_wrap_obj(object_name_with_field,False,[],[],cnd_obj_set[child_obj[1]])
                        except KeyError:
                            child_obj_check = child_wrap_obj(object_name_with_field,False,[],[],'')
                    else:
                        child_obj_check = child_wrap_obj(object_name_with_field,False,[],[],'')
                    child_wrapper = generate_child_obj(child_obj,child_obj_check)
                    child_obj_list.append(child_obj[1])
                    child_wrapper_list.append(child_wrapper)
                else:
                    child_wrapper = generate_child_obj(child_obj,child_wrapper_list[child_obj_list.index(child_obj[1])])
                    child_wrapper_list[child_obj_list.index(child_obj[1])] = child_wrapper
                old_child_obj_meta.append(child_obj)
        if len(field_list) > 0 :
            obj_wrapper.childObjWrapperList = child_wrapper_list
        else:
            obj_wrapper = obj_wrap(main_obj,False,[],[],[])
            obj_wrapper.childObjWrapperList = child_wrapper_list
        obj_wrapper = json.dumps(obj_wrapper, default=lambda o: o.__dict__)
        print("ObjMetaDataInfo-->{}".format(obj_wrapper))
        data = {
            "recordId" : record_id,
            "jsonData" : obj_wrapper,
            "folderId": folder_id_dyn,
            "fileName" : file_name
        }
        json_dic = json.loads(obj_wrapper)
        collection_name = db[json_dic.get('objName')]
        formed_query = generate_mongodb_query(json_dic,record_id)
        queried_data = collection_name.aggregate(formed_query)
        record_data = list(queried_data)
        new_data = record_data[0]
        print('queried_data-->',record_data[0])
        for obj_field in new_data:
            if type(new_data[obj_field]) in (tuple, list) :
               new_value = {'records' : new_data[obj_field]}
               new_data[obj_field] = new_value
        bind_values_doc(new_data,doc)
        doc.save('test.docx')
        return 'Success'
    else :
        obj_wrapper = json.dumps(obj_wrapper, default=lambda o: o.__dict__)
        print("ObjMetaDataInfo-->{}".format(obj_wrapper))
        data = {
            "recordId" : record_id,
            "jsonData" : obj_wrapper,
            "folderId": folder_id_dyn,
            "fileName" : file_name
        }
        json_dic = json.loads(obj_wrapper)
        collection_name = db[json_dic.get('objName')]
        formed_query = generate_mongodb_query(json_dic,record_id)
        queried_data = collection_name.aggregate(formed_query)
        record_data = list(queried_data)
        print('queried_data-->',record_data[0])
        new_data = record_data[0]
        for obj_field in new_data:
            if type(new_data[obj_field]) in (tuple, list) :
               new_value = {'records' : new_data[obj_field]}
               new_data[obj_field] = new_value
        bind_values_doc(new_data,doc)
        doc.save('testdocx.docx')
        return 'Success'

def get_all_table_patterns(whole_text):
    table_patterns = re.findall("\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:.*?\\}",whole_text)
    if len(table_patterns) > 0:
        table_pattern_list.append(table_patterns[0])
        remaining_text = whole_text.index('END')
        print(remaining_text,whole_text[remaining_text+3:])
        if remaining_text != -1 :
            return get_all_table_patterns(whole_text[remaining_text+3:])
        else :
            return table_pattern_list
    else :
        return table_pattern_list

#generate dynamic query in mongodb
def generate_mongodb_query(json_dic,record_id):
    mongodbquery = []
    final_projection_list = []
    parent_object_list = []
    print("ObjectId-->",record_id)
    objMatch = {
            '$match':{
                '_id' : ObjectId(record_id)
            }
        }
   
    mongodbquery.append(objMatch)  
    final_projection_list.append('_id') 
    for fields_in_mainobject in json_dic.get('fieldWrapperList'):
        field_api_name = ''
        print('fields_in_mainobject-->',fields_in_mainobject)
        main_obj_api = fields_in_mainobject['fieldName'].split('(')
        field_api_name = main_obj_api[0]
        
        final_projection_list.append(fields_in_mainobject['fieldName'])
        pass
    if len(json_dic['parentObjWrapperList']) > 0 :
        for parent_obj_list in json_dic.get('parentObjWrapperList'):
            #  get field name and api name from the parent object
            parent_obj_api = parent_obj_list.get('objName').split('(')
            obj_name = parent_obj_api[1].split(')')[0]
            field_name = parent_obj_api[0]
            field_query_list = []
            # get all fields in the variable
            for fields in parent_obj_list.get('fieldWrapperList'):
                field_query_list.append(fields.get('fieldName'))
            field_query = Convert(field_query_list)
            parent_object_list.append(parent_obj_list.get('objName'))
            parent_obj = {
                "$lookup":{
                    "from": obj_name,
                    "let":{
                            "id": '$'+field_name
                         },
                    "pipeline":[
                        {
                        "$match":{
                            "$expr":{
                                "$eq":[
                                    "$_id",
                                    "$$id"
                                ]
                            }
                        }
                      },
                                {
                        "$project": field_query
                        }
                    ],
                        "as":parent_obj_list.get('objName')
                    }
                }
            mongodbquery.append(parent_obj)
    if len(json_dic.get('childObjWrapperList')) > 0 :
        for child_obj_value in json_dic.get('childObjWrapperList'):
            child_obj_api = child_obj_value.get('objName').split('(')
            field_parent_name = child_obj_api[1].split(')')[0]
            obj_child_name = child_obj_api[0]
            field_query_list = []
            # get all fields in the variable
            for fields in child_obj_value.get('fieldWrapperList'):
                field_query_list.append(fields.get('fieldName'))
            field_query = Convert(field_query_list)
            
            parent_obj = {
                "$lookup":{
                    "from": obj_child_name,
                    "let":{
                            "id": "$_id"
                         },
                    "pipeline":[
                        {
                        "$match":{
                            "$expr":{
                                "$eq":[
                                   '$'+field_parent_name,
                                    "$$id"
                                ]
                            }
                        }
                      },
                                {
                        "$project": field_query
                        }
                    ],
                        "as":child_obj_value.get('objName')
                    }
                }
            mongodbquery.append(parent_obj) 
            final_projection_list.append(child_obj_value.get('objName'))
            
    
    all_object_projection = Convert(final_projection_list)
    for key,value in all_object_projection.items():
        if key in parent_object_list:
           all_object_projection[key] = {
                                        "$arrayElemAt":[
                                        "$"+key,
                                        0
                                        ]
                                        }
          
    final_project = {
        "$project" : all_object_projection
    }
    mongodbquery.append(final_project)
    return mongodbquery

def Convert(lst): 
    res_dct = {lst[i]: 1 for i in range(0, len(lst), 1)} 
    return res_dct


#To bind values to the fields which are not inside the table in the document
def bind_values_doc(data_dict,doc):
    # doc = docx.Document(file_path)
    # docume = Document(file_path)
    full_text_after = []
    child_obj_metadata = []
    for para in doc.paragraphs:
        full_text_after.append(para.text)
    document_data_after = '\n'.join(full_text_after)
    print('document_data_after-->',document_data_after)
    withouttable = get_all_table_patterns(document_data_after.replace('\n', ' ').replace('\r', ''))
    if len(withouttable) > 0 : 
        child_obj_metadata = re.findall("\\$\\{(.*?)\\}", withouttable[0])
    for paragraph in doc.paragraphs:
        adjust_pattern  = re.findall("\\{{ADJUST\\:(.*?)\\}}", paragraph.text)
        if len(adjust_pattern) > 0 :
            matched_patterns = re.findall("\\$\\{(.*?)\\}", adjust_pattern[0])
            format_type = re.findall("\\((.*?)\\)",adjust_pattern[0])[0]
            format_type = format_type.split(',')
            date_value = attach_field_values(matched_patterns[0],data_dict)
            separate_date = date_value.split('-')
            datefield = separate_date[2][:2]
            value = date(int(separate_date[0]), int(separate_date[1]), int(datefield))
            value = value + datetime.timedelta(int(format_type[1])*365/12)
            value = value + datetime.timedelta(days=int(format_type[0]))
            value = value + datetime.timedelta(int(format_type[2])*365)
            paragraph.text = str(value)
            target_stream = StringIO()
        matched_patterns = re.findall("\\$\\{(.*?)\\}", paragraph.text)
        function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", paragraph.text)
        if len(matched_patterns) > 0 and matched_patterns[0] not in child_obj_metadata:
            for value in matched_patterns :
                            text_in_cell = paragraph.text
                            field_value = attach_field_values(value,data_dict)
                            field_value = text_in_cell.replace('${'+value+'}',str(field_value))
                            paragraph.text = field_value
                            target_stream = StringIO()
        if len(function_list) > 0 :
            field_value = ''
            field_value = generate_functions(function_list,data_dict)
            paragraph.text = field_value
        target_stream = StringIO()
        # doc.save(doc)
    
    para_table_obj = []
    child_tbl_objs = set()
     #Bind values outside child table
    for paragraph in doc.paragraphs:
        child_obj_metadata = re.findall("\\$\\{(.*?)\\}", paragraph.text)
        if len(child_obj_metadata) > 0 :
            child_obj_fields = child_obj_metadata[0].split('.')
            child_tbl_objs.add(child_obj_fields[1])
    
    #Bind SUM values outside the cild table
    for cell in doc.paragraphs:
        count_func_list = re.findall("\\{\\{RowCount:(.*?)\\}}", cell.text)
        has_sum_func = re.findall("\\SUM\\{(.*?)\\}", cell.text)
        if len(count_func_list) > 0 :
            for value in count_func_list :
                    text_in_cell = cell.text
                    field_value = str(len(data_dict[count_func_list[0]]['records']))
                    field_value = text_in_cell.replace('{{RowCount:'+value+'}}',field_value)
                    cell.text = field_value
        elif len(has_sum_func) > 0 :
            format_type = re.findall("(#[A-Z]*)",has_sum_func[0])
            splited_list = has_sum_func[0].split('.')
            if len(format_type) > 0 :
                corrected_field = ''
                if len(format_type) > 0 :
                    corrected_field = splited_list[-1].replace(format_type[0],'').rstrip()
                else :
                    corrected_field = splited_list[-1]
                sum_of_field = 0 
                formatted_type_data = ''
                for field in data_dict[splited_list[1]]['records'] :
                    if len(splited_list) > 0 :
                        if len(splited_list) == 3 :
                            try :
                                formatted_type_data = str(field[corrected_field])
                            except:
                                formatted_type_data = ''
                            formatted_type = str(formatted_type_data)
                        elif len(splited_list) == 4:
                            obj_name_match = re.split('Id',splited_list[2])
                            try :
                                formatted_type_data = field[obj_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]].keys() else ''
                            except : 
                                formatted_type_data = ''
                            formatted_type = str(formatted_type_data)
                        elif len(splited_list) == 5 :
                            obj_name_match = re.split('Id',splited_list[2])
                            field_name_match = re.split('Id',splited_list[3])
                            try :
                                formatted_type_data = field[obj_name_match[0]][field_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]][field_name_match].keys() else ''
                            except :
                                formatted_type_data = ''
                            formatted_type = str(formatted_type_data)
                    print(formatted_type_data,type(formatted_type_data))
                    sum_of_field = sum_of_field + float(formatted_type_data)
                    # sum_of_field = sum_of_field + float(field[corrected_field])
                text_in_cell = cell.text
                curr_value= locale.currency(sum_of_field, grouping=True)
                value = curr_value
                field_value = text_in_cell.replace('$SUM{'+has_sum_func[0]+'}',value)
                cell.text = field_value
            else:
                corrected_field = ''
                if len(format_type) > 0 :
                    corrected_field = splited_list[-1].replace(format_type[0],'').rstrip()
                else :
                    corrected_field = splited_list[-1]
                sum_of_field = 0 
                formatted_type_data = ''
                for field in data_dict[splited_list[1]]['records'] :
                    if len(splited_list) > 0 :
                        if len(splited_list) == 3 :
                            formatted_type = str(field[corrected_field])
                        elif len(splited_list) == 4:
                            obj_name_match = re.split('Id',splited_list[2])
                            formatted_type_data = field[obj_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]].keys() else ''
                            formatted_type = str(formatted_type_data)
                        elif len(splited_list) == 5 :
                            obj_name_match = re.split('Id',splited_list[2])
                            field_name_match = re.split('Id',splited_list[3])
                            try :
                                formatted_type_data = field[obj_name_match[0]][field_name_match][corrected_field] if corrected_field in field[obj_name_match[0]][field_name_match].keys() else ''
                            except :
                                formatted_type_data = ''
                            formatted_type = str(formatted_type_data)
                    sum_of_field = sum_of_field + float(formatted_type)
                text_in_cell = cell.text
                field_value = str(sum_of_field)
                field_value = text_in_cell.replace('$SUM{'+has_sum_func[0]+'}',field_value)
                cell.text = field_value
    table_obj_to_bind_list = []
    for objects in child_tbl_objs :
         fields_list = []
         for paragraph in doc.paragraphs:
             child_obj_metadata = re.findall("\\$\\{(.*?)\\}", paragraph.text)
             if len(child_obj_metadata) > 0 :
                 child_obj_fields = child_obj_metadata[0].split('.')
                 if objects == child_obj_fields[1] :
                     fields_list.append(child_obj_metadata[0])
         table_obj_to_bind_list.append({'objName':objects,'fieldList':fields_list})
    

    for just_iterate in table_obj_to_bind_list :
        for paragraph in doc.paragraphs:
            table_obj = re.findall("\\$tbl{START:(.*):", paragraph.text)
            if len(table_obj) == 0:
                table_obj = re.findall("\\$tbl\\{START:(.*?)\\}", paragraph.text)
            if len(table_obj) > 0 :
                if just_iterate['objName'] == table_obj[0]:
                    for record in data_dict[table_obj[0]]['records'] :
                        for fields in just_iterate['fieldList'] :
                            child_obj_fields = fields.split('.')
                            field_pattern = fields.split('.')
                            format_type = re.findall("(#[A-Z]*)",fields)
                            corrected_field = ''
                            if len(format_type) > 0 :
                                corrected_field = field_pattern[-1].replace(format_type[0],'').rstrip()
                            else :
                                corrected_field = field_pattern[-1]
                                splited_list = corrected_field.split('.')
                            if len(field_pattern) == 3 :
                                try :
                                    formatted_type_data = record[corrected_field]
                                except :
                                    formatted_type_data = ''
                                formatted_type = str(formatted_type_data)
                            elif len(field_pattern) == 4:
                                obj_name_match = re.split('Id',field_pattern[2])
                                try :
                                    formatted_type_data = record[obj_name_match[0]][corrected_field] if corrected_field in record[obj_name_match[0]].keys() else ''
                                except :
                                    formatted_type_data = ''
                                formatted_type = str(formatted_type_data)
                            elif len(field_pattern) == 5 :
                                obj_name_match = re.split('Id',field_pattern[2])
                                field_name_match = re.split('Id',field_pattern[3])
                                try :
                                    formatted_type_data = record[obj_name_match[0]][field_name_match[0]][corrected_field] if corrected_field in record[obj_name_match[0]][field_name_match[0]].keys() else ''
                                except :
                                    formatted_type_data = ''
                                formatted_type = str(formatted_type_data)
                            if len(format_type) > 0 and format_type[0] == '#NUMBER' :
                                value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                            elif len(format_type) > 0 and format_type[0] == '#CURRENCY' :
                                    curr_value= locale.currency(formatted_type_data, grouping=True)
                                    value = curr_value
                            elif len(format_type) > 0 and format_type[0] == '#DATE' :
                                separate_date = formatted_type.split('-')
                                datefield = separate_date[2][:2]
                                value = date(int(separate_date[0]), int(separate_date[1]), int(datefield)).ctime()
                                value = value.split(' ')
                                value = value[1]+' '+value[2]+','+''+value[-1]
                            field_name = value if len(format_type) > 0 else formatted_type
                            paragraph.insert_paragraph_before(field_name)
                            target_stream = StringIO()

    for just_iterate in table_obj_to_bind_list :
        for paragraph in doc.paragraphs:
            table_obj = re.findall("\\$tbl{START:(.*):", paragraph.text)
            table_end = re.findall("\\$tbl\\{END:(.*?)\\}", paragraph.text)
            if len(table_obj) == 0 :
                table_obj = re.findall("\\$tbl\\{START:(.*?)\\}", paragraph.text)
            if len(table_obj) > 0 :
                if just_iterate['objName'] == table_obj[0]:
                   paragraph.text = ''
            elif len(table_end) > 0 :
                if just_iterate['objName'] == table_end[0]:
                    paragraph.text = ''
            else :
                child_obj_metadata = re.findall("\\$\\{(.*?)\\}", paragraph.text)
                if len(child_obj_metadata) > 0 :
                    child_obj_fields = child_obj_metadata[0].split('.')
                    if just_iterate['objName'] == child_obj_fields[1] :
                        paragraph.text = ''
    target_stream = StringIO()


    alltext_in_tbl = []
    for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    alltext_in_tbl.append(cell.text)
    alltext_in_tbl = '\n'.join(alltext_in_tbl)
    table_patterns = re.findall("\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:.*?\\}", alltext_in_tbl.replace('\n', ' ').replace('\r', ''))
    table_pattern_list = get_all_table_patterns(alltext_in_tbl.replace('\n', ' ').replace('\r', ''))
    table_pattern_string = ' '.join(table_pattern_list)
    for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matched_patterns = re.findall("\\$\\{(.*?)\\}", cell.text)
                    function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", cell.text)
                    field_value = ''
                    if len(function_list) > 0 :
                            field_value = generate_functions(function_list,data_dict)
                            cell.text = field_value
                    elif len(matched_patterns) > 0 :
                        for value in matched_patterns :
                            if len(table_patterns)>0 and matched_patterns[0] in table_pattern_string:
                                text_in_cell = cell.text
                                field_value = attach_field_values(value,data_dict)
                                field_value = text_in_cell.replace('${'+value+'}',str(field_value))
                                pass
                            else:
                                text_in_cell = cell.text
                                field_value = attach_field_values(value,data_dict)
                                field_value = text_in_cell.replace('${'+value+'}',str(field_value))
                                # cell.text = field_value
    
    target_stream = StringIO()
    
    # Iterating tables to bind parent field values
    if len(doc.tables) > 0:
        table_fields_list = []
        alltext_in_doc =[]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    alltext_in_doc.append(cell.text)
                    # cell.text = attach_field_values(cell.text,data_dict,file_path)
                    matched_patterns = re.findall("\\$\\{(.*?)\\}", cell.text)
                    function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", cell.text)
                    
                    if len(matched_patterns) > 0 :
                        if len(table_patterns)>0 and matched_patterns[0] in table_pattern_string:
                                pass
                        else:
                            for value in matched_patterns :
                                    cell.text = attach_field_values(value,data_dict)
                    elif len(function_list) > 0 :
                        field_value = ''
                        field_value = generate_functions(function_list,data_dict)
                        cell.text = field_value
                    target_stream = StringIO()
                    
                    # doc.save(doc)
        alltext_in_doc = '\n'.join(alltext_in_doc)
        table_values = re.findall("\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:.*?\\}", alltext_in_doc.replace('\n', ' ').replace('\r', ''))
        
        if len(table_values) > 0 :
            table_fields_list = re.findall("\\$\\{(.*?)\\}",table_pattern_string)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        table_row_details = re.findall("\\$\\{(.*?)\\}", cell.text)
                        count_func_list = re.findall("\\{\\{RowCount:(.*?)\\}}", cell.text)
                        has_sum_func = re.findall("\\SUM\\{(.*?)\\}", cell.text)
                        if len(count_func_list) > 0 :
                            for value in count_func_list :
                                    text_in_cell = cell.text
                                    field_value = str(len(data_dict[count_func_list[0]]['records']))
                                    field_value = text_in_cell.replace('{{RowCount:'+value+'}}',field_value)
                                    cell.text = field_value
                        elif len(has_sum_func) > 0 :
                            format_type = re.findall("(#[A-Z]*)",has_sum_func[0])
                            splited_list = has_sum_func[0].split('.')
                            if len(format_type) > 0 :
                                corrected_field = ''
                                if len(format_type) > 0 :
                                    corrected_field = splited_list[-1].replace(format_type[0],'').rstrip()
                                else :
                                    corrected_field = splited_list[-1]
                                sum_of_field = 0 
                                formatted_type_data = ''
                                for field in data_dict[splited_list[1]]['records'] :
                                    if len(splited_list) > 0 :
                                        if len(splited_list) == 3 :
                                            formatted_type = str(field[corrected_field])
                                            formatted_type_data = field[corrected_field]
                                        elif len(splited_list) == 4:
                                            obj_name_match = re.split('Id',splited_list[2])
                                            formatted_type_data = field[obj_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]].keys() else ''
                                            formatted_type = str(formatted_type_data)
                                        elif len(splited_list) == 5 :
                                            obj_name_match = re.split('Id',splited_list[2])
                                            field_name_match = re.split('Id',splited_list[3])
                                            try :
                                                formatted_type_data = field[obj_name_match[0]][field_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]][field_name_match].keys() else ''
                                            except :
                                                formatted_type_data = ''
                                            formatted_type = str(formatted_type_data)
                                    sum_of_field = sum_of_field + float(formatted_type_data)
                                    # sum_of_field = sum_of_field + float(field[corrected_field])
                                text_in_cell = cell.text
                                curr_value= locale.currency(sum_of_field, grouping=True)
                                value = curr_value
                                field_value = text_in_cell.replace('$SUM{'+has_sum_func[0]+'}',value)
                                cell.text = field_value
                            else:
                                corrected_field = ''
                                if len(format_type) > 0 :
                                    corrected_field = splited_list[-1].replace(format_type[0],'').rstrip()
                                else :
                                    corrected_field = splited_list[-1]
                                sum_of_field = 0 
                                formatted_type_data = ''
                                for field in data_dict[splited_list[1]]['records'] :
                                    if len(splited_list) > 0 :
                                        if len(splited_list) == 3 :
                                            try :
                                                formatted_type_data = str(field[corrected_field])
                                            except :
                                                formatted_type_data = ''
                                            formatted_type = str(formatted_type_data)
                                        elif len(splited_list) == 4:
                                            obj_name_match = re.split('Id',splited_list[2])
                                            try :
                                                formatted_type_data = field[obj_name_match[0]][corrected_field] if corrected_field in field[obj_name_match[0]].keys() else ''
                                            except :
                                                formatted_type_data = ''
                                            formatted_type = str(formatted_type_data)
                                        elif len(splited_list) == 5 :
                                            obj_name_match = re.split('Id',splited_list[2])
                                            field_name_match = re.split('Id',splited_list[3])
                                            try :
                                                formatted_type_data = field[obj_name_match[0]][field_name_match][corrected_field] if corrected_field in field[obj_name_match[0]][field_name_match].keys() else ''
                                            except :
                                                formatted_type_data = ''
                                            formatted_type = str(formatted_type_data)
                                    sum_of_field = sum_of_field + float(formatted_type)
                                text_in_cell = cell.text
                                field_value = str(sum_of_field)
                                field_value = text_in_cell.replace('$SUM{'+has_sum_func[0]+'}',field_value)
                                cell.text = field_value
                        
                        
                        if len(table_row_details) > 0 and table_row_details[0] not in table_fields_list :
                            matched_patterns = re.findall("\\$\\{(.*?)\\}", cell.text)
                            function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", cell.text)
                            field_value = ''
                            if len(function_list) > 0 :
                                    field_value = generate_functions(function_list,data_dict)
                                    cell.text = field_value
                            elif len(matched_patterns) > 0 :
                                table_obj = re.findall("\\$tbl\\{START:(.*?)\\}", cell.text)
                                for value in matched_patterns :
                                        text_in_cell = cell.text
                                        field_value = attach_field_values(value,data_dict)
                                        field_value = text_in_cell.replace('${'+value+'}',field_value)
                                        cell.text = field_value
            target_stream = StringIO()
            # r = requests.post("https://yourInstance.salesforce.com/services/data/v23.0/sobjects/ContentVersion",data=obj_wrapper,headers = 
            # curl https://yourInstance.salesforce.com/services/data/v23.0/sobjects/ContentVersion -H "Authorization: Bearer token" -H "Content-Type: multipart/form-data; boundary=\"boundary_string\"" --data-binary @NewContentVersion.json
            # doc.save(doc)
 
        def remove_row(table, row):
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)          
        
        # Iterating tables to bind child field values       
        for table in doc.tables:
            column_value_list = []
            head_obj = []
            row_to_add = []
            for row_index,row in enumerate(table.rows) : 
                    for column_index,cell in enumerate(row.cells):
                            check_child =  re.findall("\\$\\{(.*?)\\}", cell.text)
                            if len(check_child) > 0 and  check_child[0] in  table_fields_list :
                                table_row_details = re.findall("\\$\\{(.*?)\\}", cell.text)
                            table_obj = re.findall("\\$tbl{START:(.*):", cell.text)
                            if len(table_obj) == 0:
                                table_obj = re.findall("\\$tbl\\{START:(.*?)\\}", cell.text)
                            table_end = re.findall("\\$tbl\\{END:(.*?)\\}", cell.text)
                            if len(table_obj) > 0:
                                row_to_add = table.row_cells(row_index)
                                head_obj = re.findall("\\$tbl{START:(.*):", cell.text)
                                if len(head_obj) == 0:
                                    head_obj = re.findall("\\$tbl\\{START:(.*?)\\}", cell.text)
                                head_obj[0] = head_obj[0].strip()
                                row_columns =[]
                                for cell in row_to_add:
                                    if cell.text not in row_columns :
                                        row_columns.append(cell.text)
                                if len(head_obj) > 0 :  
                                    if len(check_child) > 0 and  check_child[0] in  table_fields_list : 
                                        for i,record in enumerate(data_dict[head_obj[0]]['records']) : 
                                            current_row = table.rows[row_index]
                                            border_copied = copy.deepcopy(current_row._tr)
                                            tr = border_copied
                                            current_row._tr.addnext(tr)
                                            for j,column in enumerate(row_columns):
                                                table_pattern  = re.findall("\\$\\{(.*?)\\}", column)
                                                if len(table_pattern) > 0 :
                                                    field_pattern = table_pattern[0].split('.')
                                                    format_type = re.findall("(#[A-Z]*)",table_pattern[0])
                                                    corrected_field = ''
                                                    if len(format_type) > 0 :
                                                        corrected_field = field_pattern[-1].replace(format_type[0],'').rstrip()
                                                    else :
                                                        corrected_field = field_pattern[-1]
                                                        splited_list = corrected_field.split('.')
                                                    if len(field_pattern) == 3 :
                                                        try : 
                                                            formatted_type_data = record[corrected_field]
                                                        except : 
                                                            formatted_type_data = ''
                                                        formatted_type = str(formatted_type_data)
                                                    elif len(field_pattern) == 4:
                                                        obj_name_match = re.split('Id',field_pattern[2])
                                                        try :
                                                            formatted_type_data = record[obj_name_match[0]][corrected_field] if corrected_field in record[obj_name_match[0]].keys() else ''
                                                        except :
                                                            formatted_type_data = ''
                                                        formatted_type = str(formatted_type_data)
                                                    elif len(field_pattern) == 5 :
                                                        obj_name_match = re.split('Id',field_pattern[2])
                                                        field_name_match = re.split('Id',field_pattern[3])
                                                        try :
                                                            formatted_type_data = record[obj_name_match[0]][field_name_match[0]][corrected_field] if corrected_field in record[obj_name_match[0]][field_name_match[0]].keys() else ''
                                                        except :
                                                            formatted_type_data = ''
                                                        formatted_type = str(formatted_type_data)
                                                    if len(format_type) > 0 and format_type[0] == '#NUMBER' :
                                                        value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                                                    elif len(format_type) > 0 and format_type[0] == '#CURRENCY' :
                                                            curr_value= locale.currency(formatted_type_data, grouping=True)
                                                            value = curr_value
                                                    elif len(format_type) > 0 and format_type[0] == '#DATE' :
                                                        separate_date = formatted_type.split('-')
                                                        datefield = separate_date[2][:2]
                                                        value = date(int(separate_date[0]), int(separate_date[1]), int(datefield)).ctime()
                                                        value = value.split(' ')
                                                        value = value[1]+' '+value[2]+','+''+value[-1]
                                                    field_name = value if len(format_type) > 0 else formatted_type
                                                    table.cell(row_index+1, j).text = field_name
                                                    # table.rows[row_index+1].height_rule = WD_ROW_HEIGHT.AUTO
                                                    # print("table-->{}".format(table.rows[row_index+1].height_rule))
                                                    table.rows[row_index+1].height = 1        
                            if len(table_end) > 0 :
                                remove_row(table, table.rows[row_index])
        target_stream = StringIO()
        # doc.save(doc) 
    
#Method to get field index
#Parameters (fieldName, metaData, objName)
def get_field_index(field_name, data,list_name,obj_name):
        obj_list = list()
        for record in data[list_name]:
            obj_list.append(record[obj_name])
        if len(obj_list) > 0:
            try:
                return obj_list.index(field_name)
            except ValueError:
                return -1
        else: 
            return -1     

#Method to manipulate functions in the document
def generate_functions(function_list,data_dict) :
    if_condition_list = re.findall("IF\\((.*?)\\)", function_list[0])
    if len(if_condition_list) > 0 :
        conditon_value,true_value,false_value = if_condition_list[0].split(',')[0],if_condition_list[0].split(',')[1],if_condition_list[0].split(',')[2]
        field_name_list = re.findall("\\$\\{(.*?)\\}", conditon_value)
        if '==' in str(conditon_value):
            conv_value_to_str = re.split('== ',str(conditon_value))
        if '!=' in str(conditon_value):
            conv_value_to_str = re.split('!= ',str(conditon_value))
        if '>=' in str(conditon_value):
            conv_value_to_str = re.split('>= ',str(conditon_value))
        if '<=' in str(conditon_value):
            conv_value_to_str = re.split('<= ',str(conditon_value))
        if '>' in str(conditon_value):
             conv_value_to_str = re.split('> ',str(conditon_value))
        if '<' in str(conditon_value):
            conv_value_to_str = re.split('< ',str(conditon_value))
        # print("conv_value_to_str-->",conditon_value)
        added_changes = conditon_value.replace(conv_value_to_str[-1],"'"+conv_value_to_str[-1]+"'")
        if len(field_name_list) > 0 :
                splited_list = field_name_list[0].split('.')
                if len(splited_list) == 2 :
                    field_value = data_dict[splited_list[1]]
                elif len(splited_list) == 3 :
                    obj_name_match = re.split('Id',splited_list[1])
                    field_value = data_dict[obj_name_match[0]][splited_list[2]]
                elif len(splited_list) == 4 :
                    parent_name_match = re.split('Id',splited_list[1])
                    grand_name_match = re.split('Id',splited_list[2])
                    field_value = data_dict[parent_name_match[0]][grand_name_match[0]][splited_list[3]]
        field_value = str(field_value)
        # field_value = field_value.replace(" ","")
        val = added_changes.replace('${'+field_name_list[0]+'}',"'"+field_value.strip()+"'")
        
        print("evalBefore-->","true_value if "+val+" else false_value")
        print("regex",bool(re.match('^(?=.*[a-zA-Z])',val)))
        if bool(re.match('^(?=.*[a-zA-Z])',val)) == False:
            val = val.replace("'","")    
        cons = eval("true_value if "+val+" else false_value")
        print("Result-->",cons)
        return cons
    else : 
        return "Error"


#To bind values from salesforce to the matched string
#Parameters(fieldName, metaData, filePath)
def attach_field_values(obj_to_bind,data_dict) :
     function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", obj_to_bind)
     field_name = ''
     if len(function_list) > 0 :
            field_name = generate_functions(function_list,data_dict)
     else :
            format_type = re.findall("(#[A-Z]*)",obj_to_bind)
            corrected_field = ''
            if len(format_type) > 0 :
                corrected_field = obj_to_bind.replace(format_type[0],'').rstrip()
            else :
                corrected_field = obj_to_bind
            splited_list = corrected_field.split('.')

            if len(splited_list) == 2 :
                formatted_type_data = data_dict[splited_list[1]] if splited_list[1] in data_dict.keys() else ''
                formatted_type = str(formatted_type_data)
                if len(format_type) > 0 and format_type[0] == '#NUMBER' :
                    value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                elif len(format_type) > 0 and format_type[0] == '#CURRENCY' :
                    curr_value= locale.currency(formatted_type_data, grouping=True)
                    value = curr_value
                    
                elif len(format_type) > 0 and format_type[0] == '#DATE' :
                    separate_date = formatted_type.split('-')
                    datefield = separate_date[2][:2]
                    value = date(int(separate_date[0]), int(separate_date[1]), int(datefield)).ctime()
                    value = value.split(' ')
                    value = value[1]+' '+value[2]+','+''+value[-1]
                field_name = value if len(format_type) > 0 else formatted_type
            elif len(splited_list) == 3 :
                obj_name_match = re.split('Id',splited_list[1])
                try :
                    formatted_type_data = data_dict[obj_name_match[0]][splited_list[2]] if splited_list[2] in data_dict[obj_name_match[0]].keys() else ''
                except KeyError:
                    formatted_type_data = ''
                formatted_type = str(formatted_type_data)
                if len(format_type) > 0 and format_type[0] == '#NUMBER' :
                    value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                elif len(format_type) > 0 and format_type[0] == '#CURRENCY' :
                    curr_value= locale.currency(formatted_type_data, grouping=True)
                    value = curr_value
                elif len(format_type) > 0 and format_type[0] == '#DATE' :
                    separate_date = formatted_type.split('-')
                    datefield = separate_date[2][:2]
                    value = date(int(separate_date[0]), int(separate_date[1]), int(datefield)).ctime()
                    value = value.split(' ')
                    value = value[1]+' '+value[2]+','+''+value[-1]
                field_name = value if len(format_type) > 0 else formatted_type
            elif len(splited_list) == 4 :
                 obj_name_match = re.split('Id',splited_list[1])
                 obj_field_name = re.split('Id',splited_list[2])
                 try :
                    formatted_type_data = data_dict[obj_name_match[0]][obj_field_name[0]][splited_list[3]]
                 except KeyError:
                     formatted_type_data = ''
                 formatted_type = str(formatted_type_data)
                #  formatted_type = data_dict[parent_name_match[0]][grand_name_match[0]][splited_list[3]] if [splited_list[3]] in parent_list.keys() else ''
                 if len(format_type) > 0 and format_type[0] == '#NUMBER' :
                    value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                 elif len(format_type) > 0 and format_type[0] == '#CURRENCY' :
                    curr_value= locale.currency(formatted_type_data, grouping=True)
                    value = curr_value
                 elif len(format_type) > 0 and format_type[0] == '#DATE' :
                    separate_date = formatted_type.split('-')
                    datefield = separate_date[2][:2]
                    value = date(int(separate_date[0]), int(separate_date[1]), int(datefield)).ctime()
                    value = value.split(' ')
                    value = value[1]+' '+value[2]+','+''+value[-1]
                 field_name = value if len(format_type) > 0 else formatted_type

     return field_name


if __name__ == "__main__":
    app.run()
    

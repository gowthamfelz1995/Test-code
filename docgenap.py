import os
from flask import Flask, request, render_template, jsonify, url_for, redirect
from oauthlib.oauth2 import WebApplicationClient
import requests 
from werkzeug.utils import secure_filename 
from xml.etree import ElementTree
from docx import Document
from docx.document import Document as _Document
import docx
from base64 import b64decode
import re
import json  
from datetime import date
import datetime
import copy
import xml.etree.ElementTree as ET
from docx.enum.table import WD_ROW_HEIGHT
from Models.objwrapper import obj_wrap
from Models.fieldwrap import field_wrap_obj
from Models.grandwrap import grand_wrap_obj
from Models.parentwrap import parent_wrap_obj
from Models.childwrap import child_wrap_obj
from Models.user import user_credentials
from _io import BytesIO, StringIO
import base64
import io
import random
import string
import locale
locale.setlocale(locale.LC_ALL, 'en_US.utf-8')
from sqlalchemy import create_engine,Column,Integer,String,Date,ForeignKey,MetaData,Table
from sqlalchemy.ext.declarative import declarative_base 
import pymysql.cursors
from _datetime import date, timedelta
import uuid
import shutil
from flask_cors import CORS
import jwt
from oauthlib.common import urlencode
from pymongo import MongoClient
client = MongoClient()



UPLOAD_FOLDER = str(os.getcwd())+'/python-mailmerge/Document'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif'}
path = str(os.getcwd())+'/Document'
app = Flask(__name__)
cors = CORS(app, resources={"/*": {"origins": "*"}})
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER 
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
file_path = ''
document_data = ''
document_data_initial = ''
record_id = ''
file_name = ''
file_type = ''
table_pattern_list = []
folder_id_dyn = ''



@app.route('/generate_file', methods =['POST'])
def generate_document():
    content = request.files['file'].read()
    record_id = '01'
    # file_name = request.files['file'].filename
    file_type = 'DOCX'
    file_name = 'test'+'.docx'
    bytes = b64decode(content)
    source_stream = BytesIO(content)
    doc = Document(source_stream)
    source_stream.close()
    full_text = []
    field_list = []
    child_obj_metadata = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    document_data = '\n'.join(full_text)
    field_list = re.findall("\\$\\{(.*?)\\}", document_data)
    child_obj_metadata = re.findall("\\$tbl\\{(.*?)\\}", document_data)
    withouttable = get_all_table_patterns(document_data.replace('\n', ' ').replace('\r', ''))
    if len(withouttable) > 0 : 
        child_obj_metadata = re.findall("\\$\\{(.*?)\\}", withouttable[0])
        field_list = list(set(field_list) - set(child_obj_metadata))
    head_child_obj = '' 
    alltext_in_tbl = []
    for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    alltext_in_tbl.append(cell.text)
    alltext_in_tbl = '\n'.join(alltext_in_tbl)
    table_patterns = re.findall("\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:[A-Za-z]*\\}", alltext_in_tbl.replace('\n', ' ').replace('\r', ''))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                fields_in_cell = re.findall("\\$\\{(.*?)\\}",cell.text)
                child_obj_values = re.findall("\\$tbl\\{(.*?)\\}",cell.text)
                child_obj_str = []
                if len(child_obj_values) > 0:
                   child_obj_str =  re.findall("\\$tbl\\{START:(.*?)\\}", cell.text) 
                if len(child_obj_str) > 0:
                   head_child_obj = child_obj_str[0]
                    
                # if len(fields_in_cell) > 0 and fields_in_cell[0].split('.')[1] == head_child_obj.strip():
                #     print("EXISTS")
                #     for field in fields_in_cell :
                #         child_obj_metadata.append(field)
                # elif len(fields_in_cell) > 0 :
                #     for field in fields_in_cell :
                #         field_list.append(field)
                if len(fields_in_cell) > 0 :
                    for field in fields_in_cell :
                        if field.split('.')[1] == head_child_obj.strip():
                            child_obj_metadata.append(field)
                        else:
                            field_list.append(field)
    field_list = list(dict.fromkeys(field_list))
    
    
    #This block will do json formation of fields and parent objects which will be returned to salesforce for retrieving data             
    if len(field_list) > 0 :
        obj_wrapper = obj_wrap(field_list[0].split('.')[0],False,[],[],[])
        field_wrapper = []
        parent_wrapper = []
        parent_field_wrapper = []
        grand_parent_field_wrapper = []
        grand_wrapper = []
        for field in field_list:
            format_type = re.findall("(#[A-Z]*)",field)
            if len(format_type) > 0 :
                corrected_field = field.replace(format_type[0].strip(),'')
            else :
                corrected_field = field
            parent_obj = corrected_field.split('.')
            if len(parent_obj) == 2:
                field_wrap = field_wrap_obj(parent_obj[-1],False)
                if field_wrap not in field_wrapper:
                    field_wrapper.append(field_wrap)

            else:
                field_wrap = field_wrap_obj(parent_obj[1],False)
                if field_wrap.__dict__ not in field_wrapper:
                    field_wrapper.append(field_wrap.__dict__)
                filtered_obj = parent_obj[1:len(parent_obj)]

                if len(filtered_obj) == 3:
                    parent_field_wrap = field_wrap_obj(filtered_obj[1],False)
                    if parent_field_wrap not in parent_field_wrapper:
                        parent_field_wrapper.append(parent_field_wrap)
                    grand_parent_field_wrap = field_wrap_obj(filtered_obj[2],False)
                    if grand_parent_field_wrap not in grand_parent_field_wrapper:
                        grand_parent_field_wrapper.append(grand_parent_field_wrap)
                    grand_wrap = grand_wrap_obj(filtered_obj[1],False,grand_parent_field_wrapper)
                    if len(grand_wrapper) > 0:
                        check_grobj_list = list()
                        for obj in grand_wrapper:
                            check_grobj_list.append(obj.objName)

                        if grand_wrap.objName not in check_grobj_list:
                            grand_wrapper.append(grand_wrap)

                        elif {
                                'fieldName': filtered_obj[2]
                        } not in grand_wrapper[check_grobj_list.index(
                                grand_wrap.objName)].fieldWrapperList:
                            grand_wrapper[check_grobj_list.index(
                                grand_wrap.objName)].fieldWrapperList.append(
                                    field_wrap_obj(filtered_obj[2],False))
                    else:
                        grand_wrapper.append(grand_wrap)
                    parent_wrap = parent_wrap_obj(filtered_obj[0],False,parent_field_wrapper,[],[],grand_wrapper)

                    if len(parent_wrapper) > 0:
                        check_obj_list = list()
                        for obj in parent_wrapper:
                            check_obj_list.append(obj.objName)
                        if parent_wrap.objName not in check_obj_list:
                            parent_wrapper.append(parent_wrap)

                        else:

                            if field_wrap_obj(filtered_obj[1],False) not in parent_wrapper[check_obj_list.index(
                                    parent_wrap.objName)].fieldWrapperList:
                                parent_wrapper[check_obj_list.index(
                                    parent_wrap.objName
                                )].fieldWrapperList.append(field_wrap_obj(filtered_obj[1],False))
                                # if 'grandObjWrapperList' in parent_wrapper[
                                #         check_obj_list.index(
                                #             filtered_obj[0])]:
                                parent_wrapper[check_obj_list.index(
                                    parent_wrap.objName
                                )].grandObjWrapperList.append(grand_wrap_obj(filtered_obj[1],False,[field_wrap_obj(filtered_obj[2],False)]))

                                # else:
                                # parent_wrapper[check_obj_list.index(
                                #     parent_wrap.objName
                                # )].grandObjWrapperList = [grand_wrap_obj(filtered_obj[1],False,[field_wrap_obj(filtered_obj[2],False)])]
                                
                                
                                
                                

                            else:

                                check_grandobj_list = list()
                                if 'grandObjWrapperList' in parent_wrapper[
                                        check_obj_list.index(
                                            parent_wrap.objName)]:

                                    for obj in parent_wrapper[check_obj_list.index(
                                            parent_wrap.objName
                                    )].grandObjWrapperList:
                                        check_grandobj_list.append(obj.objName)
                                else:
                                    check_grandobj_list = []
                                if grand_wrap.objName not in check_grandobj_list:
                                    grand_wrapper.append(grand_wrap)

                                else:

                                    if field_wrap_obj(filtered_obj[2],False) not in parent_wrapper[check_obj_list.index(
                                            parent_wrap.objName
                                    )].grandObjWrapperList[
                                            check_grandobj_list.index(
                                                grand_wrap.objName
                                            )].fieldWrapperList:

                                        parent_wrapper[check_obj_list.index(
                                            parent_wrap.objName
                                        )].grandObjWrapperList[
                                            check_grandobj_list.index(
                                                grand_wrap.objName
                                            )].fieldWrapperList.append(field_wrap_obj(filtered_obj[2],False))

                    else:
                        parent_wrapper.append(parent_wrap)

                    grand_wrap = {}
                    grand_parent_field_wrap = {}
                    parent_field_wrap = {}
                    parent_wrap = {}
                    parent_field_wrapper = []
                    grand_parent_field_wrapper = []
                    grand_wrapper = []

                elif len(filtered_obj) == 2:
                    parent_field_wrap = field_wrap_obj(filtered_obj[-1],False)
                    check_parent_obj_list = list()
                    for obj in parent_wrapper:
                        check_parent_obj_list.append(obj.objName)
                    if ({
                            'objName': filtered_obj[0]
                    } in check_parent_obj_list
                        ) and parent_field_wrap not in parent_wrapper[
                            check_parent_obj_list.index(
                                filtered_obj[0])]['fieldWrapperList']:
                        parent_field_wrapper.append(parent_field_wrap)

                    parent_wrap = parent_wrap_obj(filtered_obj[0],False,parent_field_wrapper,[],[],[])
                    if len(parent_wrapper) > 0:
                        check_obj_list = list()
                        for obj in parent_wrapper:
                            check_obj_list.append(obj.objName)
                        if parent_wrap.objName not in check_obj_list:
                            parent_wrap.fieldWrapperList = [field_wrap_obj(filtered_obj[-1],False)]
                            parent_wrapper.append(parent_wrap)
                            parent_field_wrapper = []

                        elif {
                                'fieldName': filtered_obj[-1],
                                'isExists': False
                        } not in parent_wrapper[check_obj_list.index(
                                parent_wrap.objName)].fieldWrapperList:
                            parent_wrapper[check_obj_list.index(
                                parent_wrap.objName
                            )].fieldWrapperList.append(field_wrap_obj(filtered_obj[-1],False))
                            parent_field_wrapper = []
                    else:
                        parent_wrap.fieldWrapperList.append(field_wrap_obj(filtered_obj[-1],False))
                        parent_wrapper.append(parent_wrap)
                        parent_field_wrapper = []
                parent_field_wrapper = []
        
        obj_wrapper.fieldWrapperList = field_wrapper
        obj_wrapper.parentObjWrapperList = parent_wrapper
        
        
       
        
        
    parent_wrapper = [] 
    parent_field_wrapper = []
    parent_field_wrap = {}
    parent_wrap = {}
    old_child_obj_meta = []
    
        
    
        
    
    #Method to return index of the obj 
    def check_obj_present(current_obj, child_wrapper):
        obj_list = list()
        for record in child_wrapper.parentObjWrapperList:
            obj_list.append(record.objName)
        if len(obj_list) > 0:
            try:
                return obj_list.index(current_obj)
            except ValueError:
                return -1
        else:
            return -1
    
    #Method to return index of the parent obj
    def check_grand_obj_present(current_obj, child_wrapper, parent_index):
        obj_list = list()

        for record in child_wrapper.parentObjWrapperList[parent_index].grandObjWrapperList:
            obj_list.append(record.objName)
        if len(obj_list) > 0:
            try:
                return obj_list.index(current_obj)
            except ValueError:
                return -1
        else:
            return -1
    #Method to return index of the field obj
    def check_field_obj_present(current_obj, child_wrapper):
        field_list = list()

        for record in child_wrapper.fieldWrapperList:
            field_list.append(record.fieldName)
        if len(field_list) > 0:
            try:
                return field_list.index(current_obj)
            except ValueError:
                return -1
        else:
            return -1
    
    #Method to generate childWrapper
    def generate_child_obj(child_obj, child_wrapper):
        child1_field_wrap = field_wrap_obj(child_obj[2],False)
        field_index = check_field_obj_present(child_obj[2], child_wrapper)
        if len(child_obj) == 3:
            if field_index == -1:
                    child_wrapper.fieldWrapperList.append(child1_field_wrap)
        if len(child_obj) == 4:
            if field_index == -1:
                child_wrapper.fieldWrapperList.append(child1_field_wrap)
            index_value = check_obj_present(child_obj[2], child_wrapper)
            if index_value == -1:
                child_wrapper.parentObjWrapperList.append(parent_wrap_obj(child_obj[2],False,[field_wrap_obj(child_obj[3],False)],[],[],[]))
            else:
                
                if field_wrap_obj(child_obj[3],False) not in child_wrapper.parentObjWrapperList[check_obj_present(child_obj[2],child_wrapper)].fieldWrapperList:
                    child_wrapper.parentObjWrapperList[check_obj_present(child_obj[2],child_wrapper)].fieldWrapperList.append(field_wrap_obj(child_obj[3],False))
        if len(child_obj) == 5:
            if field_index == -1:
                child_wrapper.fieldWrapperList.append(child1_field_wrap)
            index_value = check_obj_present(child_obj[2], child_wrapper)
            if index_value == -1:
                child_wrapper.parentObjWrapperList.append(parent_wrap_obj(child_obj[2],False,
                [field_wrap_obj(child_obj[3],False)],
                [grand_wrap_obj(child_obj[3],False,[field_wrap_obj(child_obj[4],False)])]
                ))
            else:
                if field_wrap_obj(child_obj[3],False) not in child_wrapper.parentObjWrapperList[index_value].fieldWrapperList:
                    child_wrapper.parentObjWrapperList[index_value].fieldWrapperList.append(field_wrap_obj(child_obj[3],False))
                grand_index_value = check_grand_obj_present(child_obj[3], child_wrapper, index_value)
                if grand_index_value == -1:
                    child_wrapper.parentObjWrapperList[index_value].grandObjWrapperList.append(grand_wrap_obj(child_obj[3],False,[field_wrap_obj(child_obj[4],False)]))
                else:
                    child_wrapper.parentObjWrapperList[index_value].grandObjWrapperList[grand_index_value].fieldWrapperList.append(field_wrap_obj(child_obj[4],False))
        return child_wrapper
    
    
            
    if len(child_obj_metadata) > 0:
        child_wrapper = child_wrap_obj('',False,[],[],'')
        child_obj_list = []
        child_wrapper_list = []
        get_condition = []
        check_whr_condition = re.findall("\\$tbl\\{END:(.*?)\\}",alltext_in_tbl)
        cnd_obj_set = dict()
        main_obj = ''
        if len(check_whr_condition) > 0 :
               for cnd_value in check_whr_condition :
                    try:
                        condition_on_table_idx = cnd_value.index('#CND')
                    except ValueError:
                        condition_on_table_idx = -1
                    if condition_on_table_idx != -1 :
                       cnd_table = cnd_value[condition_on_table_idx+4:]
                       cnd_obj = cnd_value[:condition_on_table_idx-1]
                       cnd_obj_set[cnd_obj]=cnd_table
                       
        for field in child_obj_metadata:
            format_type = re.findall("(#[A-Z]*)",field)
            if len(format_type) > 0 :
                corrected_field = field.replace(format_type[0],'')
            else :
                corrected_field = field
            child_obj = corrected_field.split('.')
            main_obj = child_obj[0]
            if child_obj not in old_child_obj_meta:
                if child_obj[1] not in child_obj_list:
                    head_obj = re.findall("\\$tbl{START:[A-Za-z]\\:(.*)",document_data)
                    if len(check_whr_condition) > 0 :
                        try :
                            child_obj_check = child_wrap_obj(child_obj[1],False,[],[],cnd_obj_set[child_obj[1]])
                        except KeyError:
                            child_obj_check = child_wrap_obj(child_obj[1],False,[],[],'')
                    else:
                        child_obj_check = child_wrap_obj(child_obj[1],False,[],[],'')
                    child_wrapper = generate_child_obj(child_obj,child_obj_check)
                    child_obj_list.append(child_obj[1])
                    child_wrapper_list.append(child_wrapper)
                else:
                    child_wrapper = generate_child_obj(child_obj,child_wrapper_list[child_obj_list.index(child_obj[1])])
                    child_wrapper_list[child_obj_list.index(child_obj[1])] = child_wrapper
                old_child_obj_meta.append(child_obj)
        if len(field_list) > 0 :
            obj_wrapper.childObjWrapperList = child_wrapper_list
        else:
            obj_wrapper = obj_wrap(main_obj,False,[],[],[])
            obj_wrapper.childObjWrapperList = child_wrapper_list
        obj_wrapper = json.dumps(obj_wrapper, default=lambda o: o.__dict__)
        print("ObjMetaDataInfo-->{}".format(obj_wrapper))
        data = {
            "recordId" : record_id,
            "jsonData" : obj_wrapper,
            "folderId": folder_id_dyn,
            "fileName" : file_name
        }
        formed_query = generate_mongodb_query(obj_wrapper)
        return json.dumps(formed_query)
    else :
        obj_wrapper = json.dumps(obj_wrapper, default=lambda o: o.__dict__)
        print("ObjMetaDataInfo-->{}".format(obj_wrapper))
        data = {
            "recordId" : record_id,
            "jsonData" : obj_wrapper,
            "folderId": folder_id_dyn,
            "fileName" : file_name
        }
        formed_query = generate_mongodb_query(obj_wrapper)
        return json.dumps(formed_query)

def get_all_table_patterns(whole_text):
    table_patterns = re.findall("\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:.*?\\}",whole_text)
    if len(table_patterns) > 0:
        table_pattern_list.append(table_patterns[0])
        remaining_text = whole_text.index('END')
        print(remaining_text,whole_text[remaining_text+3:])
        if remaining_text != -1 :
            return get_all_table_patterns(whole_text[remaining_text+3:])
        else :
            return table_pattern_list
    else :
        return table_pattern_list

#generate dynamic query in mongodb
def generate_mongodb_query(jsonData):
    mongodbquery = []
    final_projection_list = []
    parent_object_list = []
    json_dic = json.loads(jsonData)
    objMatch = {
            '$match':{
                '_id' : json_dic['objName']
            }
        }
    mongodbquery.append(objMatch)  
    final_projection_list.append('_id') 
    for fields_in_mainobject in json_dic['fieldWrapperList']:
        field_api_name = ''
        main_obj_api = fields_in_mainobject['fieldName'].split('(')
        field_api_name = main_obj_api[0]
        final_projection_list.append(field_api_name)
        pass
    if len(json_dic['parentObjWrapperList']) > 0 :
        for parent_obj_list in json_dic['parentObjWrapperList']:
            #  get field name and api name from the parent object
            parent_obj_api = parent_obj_list['objName'].split('(')
            obj_name = parent_obj_api[1].split(')')[0]
            field_name = parent_obj_api[0]
            field_query_list = []
            # get all fields in the variable
            for fields in parent_obj_list['fieldWrapperList']:
                field_query_list.append(fields['fieldName'])
            field_query = Convert(field_query_list)
            parent_object_list.append(field_name)
            parent_obj = {
                "$lookup":{
                    "from": obj_name,
                    "let":{
                            "id": '$'+field_name
                         },
                    "pipeline":[
                        {
                        "$match":{
                            "$expr":{
                                "$eq":[
                                    "$_id",
                                    "$$id"
                                ]
                            }
                        }
                      },
                                {
                        "$project": field_query
                        }
                    ],
                        "as":field_name
                    }
                }
            mongodbquery.append(parent_obj)
    if len(json_dic['childObjWrapperList']) > 0 :
        for child_obj_value in json_dic['childObjWrapperList']:
            field_query_list = []
            # get all fields in the variable
            for fields in parent_obj_list['fieldWrapperList']:
                field_query_list.append(fields['fieldName'])
            field_query = Convert(field_query_list)
            parent_obj = {
                "$lookup":{
                    "from": child_obj_value['objName'],
                    "let":{
                            "id": "$_id"
                         },
                    "pipeline":[
                        {
                        "$match":{
                            "$expr":{
                                "$eq":[
                                    json_dic['objName'],
                                    "$$id"
                                ]
                            }
                        }
                      },
                                {
                        "$project": field_query
                        }
                    ],
                        "as":child_obj_value['objName']
                    }
                }
            mongodbquery.append(parent_obj) 
            final_projection_list.append(child_obj_value['objName'])
            
    
    all_object_projection = Convert(final_projection_list)
    
    for key,value in all_object_projection.items():
        if key in parent_object_list:
           all_object_projection[key] = {
                                        "$arrayElemAt":[
                                        "$$"+key,
                                        0
                                        ]
                                        }
          
    final_project = {
        "$project" : all_object_projection
    }
    mongodbquery.append(final_project)
    print('mongodbquery-->',mongodbquery)
    return mongodbquery

def Convert(lst): 
    res_dct = {lst[i]: 1 for i in range(0, len(lst), 1)} 
    return res_dct

if __name__ == "__main__":
    app.run()
    
